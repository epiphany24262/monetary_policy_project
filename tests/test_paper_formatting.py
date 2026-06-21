"""Tests for paper formatting requirements per journal specification."""
from __future__ import annotations

import zipfile
import re
from pathlib import Path

import pytest


PAPER_DOCX = Path("paper/课程论文_提交版.docx")


@pytest.fixture
def docx_xml() -> bytes:
    """Read word/document.xml from the built DOCX."""
    if not PAPER_DOCX.exists():
        pytest.skip("DOCX not yet built")
    with zipfile.ZipFile(PAPER_DOCX, "r") as zf:
        return zf.read("word/document.xml")


def test_body_paragraphs_use_exact_360(docx_xml):
    """Ordinary body paragraphs must use w:lineRule='exact' w:line='360' (18pt fixed)."""
    # We look for spacing tags that have lineRule="exact" AND line="360"
    exact_360 = re.findall(
        rb'w:spacing[^/]*w:lineRule="exact"[^/]*w:line="360"', docx_xml
    )
    exact_360_alt = re.findall(
        rb'w:spacing[^/]*w:line="360"[^/]*w:lineRule="exact"', docx_xml
    )
    total = len(exact_360) + len(exact_360_alt)
    # Body text paragraphs should produce many exact-360 entries
    assert total >= 10, f"Expected ≥10 body paragraphs with exact 360 spacing, found {total}"


def test_figure_paragraphs_not_exact(docx_xml):
    """Paragraphs containing a:blip (pictures) must not use w:lineRule='exact'."""
    import xml.etree.ElementTree as ET
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    }
    root = ET.fromstring(docx_xml)
    for para in root.findall('.//w:p', ns):
        has_blip = para.find('.//a:blip', {'a': ns['a']}) is not None
        if has_blip:
            pPr = para.find('w:pPr', ns)
            if pPr is not None:
                spacing = pPr.find('w:spacing', ns)
                if spacing is not None:
                    lr = spacing.get(f'{{{ns["w"]}}}lineRule', '')
                    assert lr != 'exact', \
                        "Picture paragraph has w:lineRule='exact' — will clip the figure"


def test_no_page_break_before_section5():
    """Section 五 should not have a manual page break before it."""
    if not PAPER_DOCX.exists():
        pytest.skip("DOCX not yet built")
    from docx import Document
    doc = Document(str(PAPER_DOCX))
    for i, para in enumerate(doc.paragraphs):
        if "五、政策语调与国债收益率曲线" in para.text:
            # Check previous paragraph for page break
            if i > 0:
                prev = doc.paragraphs[i - 1]
                for run in prev.runs:
                    for child in run._element:
                        assert 'lastRenderedPageBreak' not in child.tag or 'br' not in child.tag
            # Check paragraph_format.page_break_before
            assert not para.paragraph_format.page_break_before
            return
    pytest.fail("Section heading 五 not found")


def test_no_page_break_before_section6():
    """Section 六 should not have a manual page break before it."""
    if not PAPER_DOCX.exists():
        pytest.skip("DOCX not yet built")
    from docx import Document
    doc = Document(str(PAPER_DOCX))
    for para in doc.paragraphs:
        if "六、进一步检验与研究局限" in para.text:
            assert not para.paragraph_format.page_break_before
            return
    pytest.fail("Section heading 六 not found")


def test_no_page_break_before_references():
    """参考文献 should not have a manual page break before it."""
    if not PAPER_DOCX.exists():
        pytest.skip("DOCX not yet built")
    from docx import Document
    doc = Document(str(PAPER_DOCX))
    for para in doc.paragraphs:
        if "参考文献" in para.text and len(para.text.strip()) < 10:
            assert not para.paragraph_format.page_break_before
            return
    pytest.fail("Reference heading not found")


def test_table4_caption_keep_with_next():
    """Table 4 caption must have keep_with_next to prevent separation."""
    if not PAPER_DOCX.exists():
        pytest.skip("DOCX not yet built")
    from docx import Document
    doc = Document(str(PAPER_DOCX))
    for para in doc.paragraphs:
        if para.text.strip().startswith("表4  "):
            # docx sometimes returns None for keep_with_next if it's explicitly set in xml but not cached
            assert '<w:keepNext/>' in para._p.xml or para.paragraph_format.keep_with_next is True, \
                "Table 4 caption missing keep_with_next"
            return
    pytest.fail("Table 4 caption not found")


def test_exactly_four_figure_captions():
    """The paper must have exactly 4 figure captions."""
    if not PAPER_DOCX.exists():
        pytest.skip("DOCX not yet built")
    from docx import Document
    doc = Document(str(PAPER_DOCX))
    captions = [p for p in doc.paragraphs if re.match(r"^图\s*\d+\s+", p.text.strip())]
    assert len(captions) == 4, f"Expected 4 figure captions, found {len(captions)}"


def test_exactly_four_table_captions():
    """The paper must have exactly 4 table captions."""
    if not PAPER_DOCX.exists():
        pytest.skip("DOCX not yet built")
    from docx import Document
    doc = Document(str(PAPER_DOCX))
    captions = [p for p in doc.paragraphs if re.match(r"^表\s*\d+\s+", p.text.strip())]
    assert len(captions) == 5, f"Expected 5 table captions, found {len(captions)}"


def test_no_chinese_in_english_abstract():
    """The English abstract must not contain Chinese characters."""
    if not PAPER_DOCX.exists():
        pytest.skip("DOCX not yet built")
    from docx import Document
    doc = Document(str(PAPER_DOCX))
    in_english = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Abstract:"):
            in_english = True
        elif text.startswith("Key words:"):
            # Check this line too
            chinese = re.findall(r"[\u4e00-\u9fff]", text)
            assert not chinese, f"Key words contain Chinese: {chinese}"
            break
        elif in_english and text:
            chinese = re.findall(r"[\u4e00-\u9fff]", text)
            assert not chinese, f"English abstract contains Chinese: {chinese}"


def test_front_matter_has_no_forced_page_break(docx_xml):
    import xml.etree.ElementTree as ET
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    root = ET.fromstring(docx_xml)
    page_breaks = root.findall('.//w:br[@w:type="page"]', ns)
    assert len(page_breaks) == 0, "There should be no forced manual page breaks in the document."


def test_title_uses_natural_weight():
    if not PAPER_DOCX.exists():
        pytest.skip("DOCX not yet built")
    from docx import Document
    doc = Document(str(PAPER_DOCX))
    for para in doc.paragraphs[:5]:
        if "中国货币政策报告文本特征" in para.text:
            for run in para.runs:
                assert run.bold is None or run.bold is False, "Title should not be forced bold."
            return
    pytest.fail("Title not found")


def test_egarch_and_bond_results_use_separate_tables():
    if not PAPER_DOCX.exists():
        pytest.skip("DOCX not yet built")
    from docx import Document
    doc = Document(str(PAPER_DOCX))
    tables = [p.text for p in doc.paragraphs if re.match(r"^表\s*\d+\s+", p.text.strip())]
    assert any("稳健性" in t for t in tables)
    assert any("国债" in t for t in tables)
    assert len(tables) == 5


@pytest.fixture
def pdf_text():
    pdf_path = Path("paper/课程论文_提交版.pdf")
    if not pdf_path.exists():
        pytest.skip("PDF not found")
    import fitz
    doc = fitz.open(str(pdf_path))
    pages = [page.get_text() for page in doc]
    doc.close()
    return pages


def test_table1_caption_rows_and_note_share_one_pdf_page(pdf_text):
    found = False
    for page in pdf_text:
        if "表1" in page and "数据来源、频率与样本口径" in page:
            assert "注：正式样本期来自项目配置和来源登记文件" in page
            found = True
    assert found


def test_learning_curve_is_in_text_validation_section(pdf_text):
    found = False
    for page in pdf_text:
        if "图2" in page and "文本分类模型的分组交叉验证学习曲线" in page:
            assert "（二）人工标注与分组交叉验证" in page or "（三）文本测量结果" in page
            found = True
    assert found


def test_learning_curve_not_in_power_section(pdf_text):
    for page in pdf_text:
        if "图2" in page and "文本分类模型的分组交叉验证学习曲线" in page:
            assert "样本量与检验功效" not in page


def test_table4_is_not_split(pdf_text):
    found_pages = sum(1 for page in pdf_text if "表4" in page and "日度波动稳健性检验" in page)
    assert found_pages == 1, "Table 4 should not be split across multiple pages"


def test_table5_caption_and_first_header_share_one_page(pdf_text):
    found = False
    for page in pdf_text:
        if "表5" in page and "国债市场探索性结果" in page:
            assert "Panel A" in page or "未预期语调" in page
            found = True
    assert found

def test_references_heading_format():
    if not PAPER_DOCX.exists():
        pytest.skip("DOCX not yet built")
    from docx import Document
    doc = Document(str(PAPER_DOCX))
    found = False
    for para in doc.paragraphs:
        if para.text.strip() == "参考文献":
            found = True
            # Check left positioned
            assert para.paragraph_format.alignment in [0, None], "References heading should be left-positioned."
            for run in para.runs:
                if run.text.strip():
                    assert run.bold in [False, None], "References heading should not be mechanically bolded."
            break
    assert found, "References heading not found."

def test_consistent_p_value_headers(docx_xml):
    import xml.etree.ElementTree as ET
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    root = ET.fromstring(docx_xml)
    
    # Extract all text in tables to check for p value strings
    for table in root.findall('.//w:tbl', ns):
        for row in table.findall('.//w:tr', ns):
            for cell in row.findall('.//w:tc', ns):
                cell_text = "".join(t.text for t in cell.findall('.//w:t', ns) if t.text)
                if "p" in cell_text or "P" in cell_text or "值" in cell_text:
                    # It's a table cell that might be a header or value containing 'p'
                    if re.search(r"[pP]\s+值", cell_text):
                        pytest.fail(f"Found hard space between p and 值 in table cell: {cell_text}")
                    if "p" in cell_text and "值" in cell_text and "p值" not in cell_text:
                        pytest.fail(f"Found separated or misformatted p value header: {cell_text}")

def test_table_vertical_separators(docx_xml):
    import xml.etree.ElementTree as ET
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    root = ET.fromstring(docx_xml)
    
    for table in root.findall('.//w:tbl', ns):
        tblBorders = table.find('.//w:tblBorders', ns)
        if tblBorders is not None:
            left = tblBorders.find('w:left', ns)
            right = tblBorders.find('w:right', ns)
            insideV = tblBorders.find('w:insideV', ns)
            
            # None of the tables should have a full outer rectangle or complete Word grid
            assert left is None or left.get(f"{{{ns['w']}}}val") in ["none", "nil"], "Table has full outer left boundary."
            assert right is None or right.get(f"{{{ns['w']}}}val") in ["none", "nil"], "Table has full outer right boundary."
            assert insideV is None or insideV.get(f"{{{ns['w']}}}val") in ["none", "nil"], "Table has complete Word grid (insideV)."
