from pathlib import Path
import re

import pandas as pd
from docx import Document



def test_paper_has_tables_figures_and_no_internal_terms():
    text = "\n".join(p.text for p in Document("paper/课程论文_提交版.docx").paragraphs)
    for banned in ["Pha" + "se", "Gate", "Co" + "dex", "P0", "P1", "P2", "JSON", "初版", "重构", "旧窗口", "旧版", "开发史", "README.md", "run_all.py", "提交目录", "缓存"]:
        assert banned not in text
    assert len(text) > 9000
    assert len([p for p in Document("paper/课程论文_提交版.docx").paragraphs if re.match(r"^图\s*\d+\s", p.text.strip())]) == 4
    assert len([p for p in Document("paper/课程论文_提交版.docx").paragraphs if re.match(r"^表\s*\d+\s", p.text.strip())]) == 5
    assert len(list(Path("output/figures").glob("journal_figure*.png"))) == 4
    assert len(list(Path("output/tables").glob("journal_table*.csv"))) == 5
    main = pd.read_csv("output/results/stock_volatility_results.csv")
    beta = f"{main.loc[main['model']=='full','beta'].iloc[0]:.4f}"
    assert beta in text
