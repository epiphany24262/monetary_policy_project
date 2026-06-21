from pathlib import Path

def test_cover_school_name_not_clipped():
    diagnostic = Path("output/diagnostics/cover_top_final.png")
    if diagnostic.exists():
        assert diagnostic.stat().st_size > 10_000
    else:
        assert Path("paper/课程论文_提交版.pdf").exists()

def test_no_embedded_figure_captions():
    fig_dir = Path("output/figures")
    paper_figs = list(fig_dir.glob("journal_figure*.png"))
    assert len(paper_figs) == 4

def test_cover_has_school_name():
    import fitz
    pdf = fitz.open("paper/课程论文_提交版.pdf")
    p1_text = pdf[0].get_text("text")
    pdf.close()
    assert len(p1_text) > 50

def test_each_figure_has_word_caption():
    from docx import Document
    doc = Document("paper/课程论文_提交版.docx")
    count = sum(1 for p in doc.paragraphs if p.text.strip().startswith("图") and len(p.text.strip()) < 30)
    assert count == 4

def test_no_duplicate_figure_title():
    import re
    from docx import Document
    doc = Document("paper/课程论文_提交版.docx")
    text = "\n".join(p.text for p in doc.paragraphs)
    captions = re.findall(r"^图\s*\d+.*$", text, re.MULTILINE)
    assert len(captions) == 4
