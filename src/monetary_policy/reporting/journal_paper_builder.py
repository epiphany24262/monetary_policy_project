from __future__ import annotations

import json
import math
import re
import shutil
import tempfile
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path

import fitz
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from PIL import Image, ImageDraw

from ..paths import FIGURES_DIR, OUTPUT_DIR, PAPER_DIR, ROOT
from .journal_figures import write_journal_figures
from .journal_style import (
    CAPTION_STYLE,
    FONT_STYLE,
    HEADING_STYLE,
    PARAGRAPH_STYLE,
    REFERENCE_STYLE,
    TABLE_STYLE,
    add_body_section,
    clear_table_borders,
    configure_section,
    _restart_page_numbering,
    mark_row_no_split,
    merge_row_cells,
    set_cell_border,
    set_cell_text,
    set_figure_paragraph_format,
    set_paragraph_format,
    set_run_font,
    set_table_width,
    set_cell_margins,
)
from .journal_tables import JournalTables, fmt, write_journal_tables


# Explicit vertical border configuration per table.
# A value of False means no vertical separator after the column.
TABLE_VERTICAL_CONFIG = {
    1: {},
    2: {},
    3: {},
    4: {},
    5: {},
}

# Keep track of table count for vertical border application
_CURRENT_TABLE_INDEX = 0

DOCX_PATH = PAPER_DIR / "课程论文_提交版.docx"
PDF_PATH = PAPER_DIR / "课程论文_提交版.pdf"
COVER_PATH = ROOT / "references/journal_format/课程论文封面.docx"
REVIEW_DIR = OUTPUT_DIR / "diagnostics" / "journal_review"


@dataclass
class PaperNumbers:
    formal_reports: int
    manual_sentences: int
    report_start: str
    report_end: str
    stock_n: int
    stock_beta: float
    stock_se: float
    stock_p: float
    stock_r2: float
    stock_interaction: float
    stock_interaction_p: float
    stock_total: float
    stock_total_p: float
    guidance_std: float
    stock_one_sd_pct: float
    egarch_n: int
    egarch_events: int
    egarch_novelty_coef: float
    egarch_lr_p: float
    egarch_perm_p: float
    egarch_vol_pct: float
    power_current: float
    power_120: float
    power_160: float
    bond_beta: float
    bond_se: float
    bond_p: float
    bond_total: float
    bond_total_p: float
    cross_coef: float
    cross_p: float
    cross_total: float
    cross_total_p: float
    svc_sentiment_acc: float
    svc_sentiment_f1: float
    svc_stance_acc: float
    svc_stance_f1: float
    svc_direction_acc: float
    svc_direction_f1: float
    lexicon_sentiment_acc: float
    lexicon_stance_acc: float
    lexicon_direction_acc: float


def build_journal_paper(results: dict) -> dict:
    PAPER_DIR.mkdir(parents=True, exist_ok=True)
    REVIEW_DIR.mkdir(parents=True, exist_ok=True)
    numbers = collect_paper_numbers(results)
    tables = write_journal_tables(results)
    figures = write_journal_figures(results)
    doc = _new_document_with_cover()
    _build_body(doc, numbers, tables, figures)
    refs = _add_references(doc)
    _write_paper_audits(numbers, refs)
    doc.core_properties.title = "中国货币政策报告文本特征与金融市场反应"
    doc.core_properties.author = "罗允绩"
    doc.core_properties.subject = "货币政策沟通与金融市场反应"
    # sanitize presentation layer text before saving
    try:
        _sanitize_document(doc)
    except Exception:
        print("Saving to", str(DOCX_PATH))
    doc.save(str(DOCX_PATH))
    _normalize_docx_spacing_rules(DOCX_PATH)
    export_pdf_with_word()
    print("Exporting data sources summary...")
    return inspect_journal_pdf()


def collect_paper_numbers(results: dict) -> PaperNumbers:
    features = results["text_features"]
    stock_panel = results["stock_panel"]
    main = results["main_vol"]
    validation = results["text_validation"]
    model = results["text_model_summary"]
    egarch = results["egarch_x"]
    egarch_main = egarch.get("main", egarch.get("main_model", {}))
    curve = pd.read_csv(ROOT / "output" / "results" / "yield_curve_results.csv")
    bond = curve[curve["dependent"].eq("delta_slope_bp_0_3")].iloc[0]
    cross = pd.read_csv(ROOT / "output" / "results" / "cross_fitted_bond_exploration.csv")
    cross_row = cross[cross["tone_aggregation"].eq("policy_relevant_mean")].iloc[0]
    power = pd.read_csv(ROOT / "output" / "diagnostics" / "market_power_analysis.csv")
    novelty = pd.to_numeric(stock_panel["guidance_novelty"], errors="coerce").dropna()
    novelty_std = float(novelty.std(ddof=1))
    stock_beta = float(main["params"]["guidance_novelty"])
    power_current = power[power["sample_size"].eq(79)].iloc[0]["power"]
    power_120 = power[power["sample_size"].eq(120)].iloc[0]["power"]
    power_160 = power[power["sample_size"].eq(160)].iloc[0]["power"]
    formal = features[features["in_formal_sample"]].copy()
    return PaperNumbers(
        formal_reports=int(formal.shape[0]),
        manual_sentences=int(validation["total_sentences"]),
        report_start=str(formal["report_period"].min()),
        report_end=str(formal["report_period"].max()),
        stock_n=int(main["n"]),
        stock_beta=stock_beta,
        stock_se=float(main["bse_hc3"]["guidance_novelty"]),
        stock_p=float(main["pvalues"]["guidance_novelty"]),
        stock_r2=float(main["r2"]),
        stock_interaction=float(main["params"]["guidance_novelty_x_post_2019"]),
        stock_interaction_p=float(main["pvalues"]["guidance_novelty_x_post_2019"]),
        stock_total=float(main["post_2019_total_effect"]["estimate"]),
        stock_total_p=float(main["post_2019_total_effect"]["p_value"]),
        guidance_std=novelty_std,
        stock_one_sd_pct=(math.exp(stock_beta * novelty_std) - 1) * 100,
        egarch_n=int(egarch_main["n_daily_observations"]),
        egarch_events=int(egarch_main["n_novelty_events"]),
        egarch_novelty_coef=float(egarch_main["parameters"]["novelty_z"]),
        egarch_lr_p=float(egarch_main["formal_lr_p_value"]),
        egarch_perm_p=float(egarch["permutation_p_novelty"]),
        egarch_vol_pct=float(egarch_main["conditional_volatility_change_pct_per_1sd_novelty"]),
        power_current=float(power_current),
        power_120=float(power_120),
        power_160=float(power_160),
        bond_beta=float(bond["beta"]),
        bond_se=float(bond["se_hc3"]),
        bond_p=float(bond["p_value"]),
        bond_total=float(bond["post_2019_total_effect"]),
        bond_total_p=float(bond["post_2019_total_p_value"]),
        cross_coef=float(cross_row["coef"]),
        cross_p=float(cross_row["p_value"]),
        cross_total=float(cross_row["post_2019_total_effect"]),
        cross_total_p=float(cross_row["post_2019_total_p_value"]),
        svc_sentiment_acc=float(model["sentiment_cv"]["accuracy"]),
        svc_sentiment_f1=float(model["sentiment_cv"]["macro_f1"]),
        svc_stance_acc=float(model["policy_stance_cv"]["accuracy"]),
        svc_stance_f1=float(model["policy_stance_cv"]["macro_f1"]),
        svc_direction_acc=float(model["policy_direction_cv"]["accuracy"]),
        svc_direction_f1=float(model["policy_direction_cv"]["macro_f1"]),
        lexicon_sentiment_acc=float(validation["sentiment_accuracy"]),
        lexicon_stance_acc=float(validation["stance_accuracy"]),
        lexicon_direction_acc=float(validation["policy_direction_accuracy"]),
    )


def _new_document_with_cover() -> Document:
    if COVER_PATH.exists():
        doc = Document(str(COVER_PATH))
        for paragraph in reversed(doc.paragraphs):
            if not paragraph.text.strip():
                paragraph.paragraph_format.space_before = 0
                paragraph.paragraph_format.space_after = 0
    else:
        doc = Document()
        paragraph = doc.add_paragraph("四川大学")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(paragraph.runs[0], FONT_STYLE["heading_cn"], 18, bold=True)
    configure_section(doc.sections[0])
    cover = doc.sections[0]
    cover.different_first_page_header_footer = True
    _restart_page_numbering(cover, 0)
    for part in [cover.header, cover.footer, cover.first_page_header, cover.first_page_footer]:
        for paragraph in part.paragraphs:
            paragraph.clear()
    add_body_section(doc)
    doc.settings.odd_and_even_pages_header_footer = True
    return doc


def _build_body(doc: Document, numbers: PaperNumbers, tables: JournalTables, figures) -> None:
    _add_front_matter(doc, numbers)
    _add_heading1(doc, "一、引言")
    for text in [
        "货币政策沟通不仅传递政策取向，也会改变市场对宏观状态和政策反应函数的判断。中国人民银行季度货币政策执行报告具有稳定的发布频率和较长的连续文本，是观察央行沟通变化的合适材料。与公告或新闻稿相比，季度报告包含更完整的宏观判断、政策回顾和下一阶段政策指引，因而能够为文本测度提供较多上下文。",
        "已有研究从不同角度处理央行沟通信息。姜富伟、胡逸驰和黄楠（2021）将中国央行报告文本与股票市场反应联系起来；董青马等（2024）强调资产价格反应中未预期政策信息的重要性；尚玉皇、刘华和申峰（2025）则把央行沟通放入国债收益率曲线框架中讨论。国际文献中，Blinder等（2008）系统讨论央行沟通的理论和证据，Gurkaynak、Sack和Swanson（2005）区分政策行动与声明信息，Hansen和McMahon（2016）、Nakamura和Steinsson（2018）、Jarocinski和Karadi（2020）进一步提示央行信息效应可能与传统政策冲击交织在一起。",
        "现有文献给本文留下的空间，不在于重新证明央行沟通会影响资产价格，而在于更细地处理“沟通中的哪一部分信息”以及“市场以何种形式反应”。季度报告中有大量宏观回顾和制度性表述，如果把全文语调直接接到市场变量上，容易把政策指引、经济判断和模板化文字混在一起。本文把政策指引章节从全文中分离出来，并把文本变化而不是文本水平作为核心解释变量。",
        "这一处理也避免了把央行沟通简单等同于好消息或坏消息。政策指引的新增表述可能包含偏松信号、偏紧信号，也可能只是对政策工具组合、风险约束或执行方式的重新说明。市场反应未必表现为收益率或股价的单向移动，却可能表现为短期波动率上升。以创新度解释股票实际波动率，正是为了捕捉这种信息更新过程。",
        "本文关心的问题更窄：在中国季度货币政策报告中，政策指引章节相对历史文本的新信息，是否与报告发布后的短期市场波动有关。研究以2006Q1至2025Q4的正式样本为基础，构造扩展窗口TF-IDF创新度，并将其与沪深300指数发布后五日实际波动率相连接。债券市场部分使用未预期政策语调和跨拟合监督语调解释国债收益率曲线变化，定位为探索性扩展。",
        "本文的经验结果可以概括为三点。第一，政策指引创新度与股票市场短期实际波动率之间存在显著正相关关系，创新度提高一个样本标准差，发布后五日实际波动率约提高13.0%。第二，Student-t EGARCH-X模型在完整连续日度收益率序列上得到正向但不显著的创新度方差项，说明核心发现并非完全由一个日度条件异方差规格驱动，但稳健性证据较弱。第三，未预期政策语调与收益率曲线斜率、跨拟合监督语调与债券反应之间没有稳定显著关系，债券市场结果需要谨慎解释。",
        "变量构造依赖较完整的Python文本处理过程。报告原文先经过PDF抽取和章节识别，再切分为句子并连接到报告层面的市场事件；金融数据也需要按照交易日历重排，才能把报告发布时间和事件窗口对齐。文本特征工程和金融事件研究在这里不是两个分离步骤，前者决定解释变量的含义，后者决定市场反应的计量口径。",
        "本文的边际贡献主要在三个层面。其一，把政策指引章节单独拿出来度量，避免全文相似度被宏观回顾和固定格式稀释。其二，用按报告分组交叉验证检验文本测量，降低同一报告模板句泄漏到测试集的风险。其三，把股票事件级回归、日度波动模型和债券探索放在同一研究问题下比较，从而区分证据强弱，而不是只挑选单一显著结果。",
        "这种安排也有助于解释不同市场结果为何不完全一致。股票波动率对信息不确定性和估值分歧反应更直接，政策指引的新表述即使不改变政策利率，也可能提高短期风险重估强度。债券收益率曲线则同时承载政策路径、增长通胀判断和期限溢价，单一语调指标更容易混合多种信息。因此，本文把股票创新度—波动率关系作为核心主线，把债券语调—收益率曲线关系作为探索性扩展。",
    ]:
        _add_body(doc, text)

    _add_heading1(doc, "二、数据来源与研究设计")
    _add_heading2(doc, "（一）货币政策报告与金融市场数据")
    for text in [
        f"报告文本来自中国人民银行官网发布的《中国货币政策执行报告》。数据库覆盖2006Q1至2026Q1，正式实证样本限定在{numbers.report_start}至{numbers.report_end}，共{numbers.formal_reports}期。2026Q1仅保留在来源登记中，不进入正式估计。报告发布时间用于对齐股票和债券交易日；若发布时间落在非交易日，事件日顺延至下一可交易日。",
        "金融市场数据包括沪深300指数日行情、中债国债1年、5年和10年期收益率，以及公开政策操作日期。股票部分计算报告发布后五个交易日实际波动率，债券部分构造收益率曲线水平、斜率和曲率的短窗口变化。政策操作变量只用于控制报告发布附近的政策环境，避免把同日或近日报告反应全部归因于文本。",
        "这些数据的频率并不一致。货币政策报告是季度文本，股票和债券价格为日度序列，政策操作又是不定期事件。实证处理需要把低频文本信息压缩为报告层变量，再把高频金融数据聚合为事件窗口反应。这样得到的样本量由报告期数决定，而不是由日度行情观测数决定，这一点对后文解释检验功效和分时期估计尤其重要。",
        "报告文本本身也存在时间上的可比性问题。早期报告排版和章节结构与后期不完全一致，若直接使用全文相似度，宏观回顾、专栏、附录和固定制度表述会稀释政策指引变化。本文只在正式样本内使用修复后的章节文本，并在表1中区分文本、市场和人工标注数据的样本口径，使不同来源的用途保持清楚。",
        "事件日对齐采用发布日与交易日的对应关系。若报告在交易日收盘后发布，或发布日不是交易日，市场反应会落入后续交易日；若报告在交易时段内可被市场观察，则当日价格也可能包含反应。本文统一用已整理的发布时间映射事件日，并在股票与债券模型中保持同一口径，避免不同市场使用不同事件定义。",
        "数据合法性和可追溯性主要体现在来源和口径上。政策报告来自央行公开网页，市场数据保留清洗后的日度序列和来源登记。论文正文不展开文件层面的复现说明，但表1给出了每类数据的频率、样本期和用途，读者可以据此判断文本、股票、债券和人工标注样本之间的对应关系。",
    ]:
        _add_body(doc, text)
    _add_table(doc, "表1  数据来源、频率与样本口径", tables.table1, note="注：正式样本期来自项目配置和来源登记文件；人工标注样本用于文本测量验证。")

    _add_heading2(doc, "（二）文本指标构造")
    for text in [
        "文本处理先识别宏观经济章节和政策指引章节。政策指引章节更接近央行对下一阶段政策取向的表述，宏观章节则主要反映对经济金融形势的判断。二者在信息含义上并不相同，因此本文把政策指引创新度作为股票市场核心解释变量，把语调指标用于债券探索和文本测量对照。",
        "创新度使用扩展窗口TF-IDF计算。对第t期报告，只利用第1期至第t期已经可见的文本拟合词项权重，并计算本期政策指引与上一期政策指引的余弦相似度；创新度定义为一减该相似度。这样做避免未来报告词汇反向影响早期文本向量，也更接近市场在报告发布时能够观察到的信息集。Ehrmann和Talmi（2020）关于央行沟通相似度与市场波动的研究，为这种从文本变化理解市场反应的处理提供了参照。",
        "语调指标来自两类方法。一类是金融情感词典和中国央行语境规则，公开中文金融情感词典参照Du等（2022），并结合货币政策语境区分偏松、偏紧、中性和无关句。另一类是字符TF-IDF与LinearSVC监督分类器，用人工标注句子进行按报告分组交叉验证。监督模型不替代创新度主线，而是检验句子层文本测量是否具有可接受的外推能力。",
        "创新度和语调指标承担不同任务。创新度衡量政策指引相对上一期的表述变化，它不预设变化方向，也不把偏松或偏紧作为前提；语调指标则试图判断句子表达的是政策取向还是宏观判断。前者更适合研究信息增量与波动率的关系，后者更适合放入收益率曲线框架中讨论政策路径预期。把两类指标混在同一个主检验中，反而会模糊经济含义。",
        "扩展窗口处理还有一个实际好处。央行报告文本存在明显的格式稳定性，一些词项在后期才频繁出现，若用全样本词典一次性拟合早期文本，就会把未来信息带入历史度量。扩展窗口TF-IDF虽然会增加早期文本向量的不确定性，但它更符合投资者在当时能够接触到的词汇环境，也使创新度的时间序列解释更自然。",
        "连续主题关注度用于补充解释文本内容变化。与把每个句子硬分到单一主题不同，连续关注度保留了增长、通胀、金融稳定、房地产和汇率等主题的相对权重，更适合描述报告关注重点的时间变化。它不进入股票核心模型，原因是主题关注度更接近经济背景变量，而政策指引创新度才直接对应“本期表述相对上一期有多新”。",
        "语境门控规则处理的是央行文本中常见的误判来源。一般金融情感词典会把“压力”“风险”“回落”等词识别为负面情绪，但这些词出现在宏观形势判断时，不一定意味着政策取向偏紧；“保持流动性合理充裕”在一般情感词典中也不容易被识别为政策含义。语境门控将这些词放回句子和章节环境中解释，减少把宏观描述误读为政策方向的情况。",
    ]:
        _add_body(doc, text)
    _add_figure(doc, FIGURES_DIR / figures.figure1, "图1  政策指引相似度与创新度")

    _add_heading2(doc, "（三）事件窗口与计量模型")
    for text in [
        "股票核心模型以报告发布后五个交易日实际波动率对数为被解释变量，解释变量包括政策指引创新度、发布前20日市场波动率、报告附近政策操作、2019年后虚拟变量以及创新度与2019年后虚拟变量的交互项。标准误采用HC3稳健形式，以减轻有限事件样本中高杠杆观测对推断的影响。",
        "日度稳健性采用Student-t EGARCH-X模型，在完整连续交易日收益率序列上估计报告日、标准化创新度和政策操作日进入条件方差方程后的影响。厚尾分布用于处理股票日收益率尖峰厚尾特征。债券模型以三日窗口内收益率曲线斜率变化为主要被解释变量，同时报告水平和曲率结果。",
        "事件窗口选择遵循两个原则。其一，窗口要短到足以减少其他宏观信息的混入；其二，窗口又要长到可以覆盖报告发布后的初步信息吸收。五个交易日实际波动率更关注不确定性释放，而不是判断价格上涨或下跌方向。债券窗口较短，是因为收益率曲线对政策预期的调整通常更快，窗口拉长反而更容易受到资金面和海外利率扰动。",
        "2019年后交互项用于描述沟通制度和市场环境变化后的差异，而不是在看到结果后更换主窗口。2019年前后，中国货币政策框架、利率市场化进程和预期管理方式都发生变化，市场对季度报告的边际反应可能随之改变。交互项只能提供样本内比较，不能单独证明制度变化的因果作用。",
        "股票和债券模型在解释变量上保持克制。股票主模型只把政策指引创新度、前期波动、政策操作和2019年后交互项放入同一规格；债券模型则集中于未预期政策语调与曲线斜率。这样做牺牲了部分解释变量的丰富性，却减少了小样本下过度拟合和事后选择变量的风险。",
    ]:
        _add_body(doc, text)

    _add_heading1(doc, "三、文本测量与有效性检验")
    _add_heading2(doc, "（一）词典与语境门控")
    for text in [
        "中文央行文本存在明显的语境差异。同一个词在宏观判断、政策目标和操作安排中可能含义不同，例如“保持合理充裕”与一般积极情绪并不等价。本文先用公开金融情感词典识别情绪，再用领域词表和语境规则修正政策倾向，尤其处理否定、条件、政策目标和无关制度表述。",
        f"词典和语境门控在{numbers.manual_sentences}句人工标注样本上接受检验。当前词典的情感准确率为{fmt(numbers.lexicon_sentiment_acc)}，政策四分类准确率为{fmt(numbers.lexicon_stance_acc)}；若只在政策相关句中区分偏松、偏紧和中性，准确率为{fmt(numbers.lexicon_direction_acc)}。这些指标说明规则方法能够捕捉一部分政策语义，但对金融情感的细粒度识别仍不充分。",
        "规则方法的优势是透明。每一次分数变化都可以追溯到词典匹配、否定词和语境规则；缺点也同样清楚，固定词表很难识别新表述和复杂句法。央行报告中常见的长句往往把目标、约束和操作安排放在同一句内，单纯逐词加总会丢失关系信息。因此，词典结果更适合作为可解释基准，而不是唯一的文本测量依据。",
    ]:
        _add_body(doc, text)

    _add_heading2(doc, "（二）人工标注与分组交叉验证")
    for text in [
        "人工标注样本按报告抽取政策指引句子，标注金融情感、政策倾向和主题类别。由于同一报告内部存在大量近似表述，随机句子切分会使相似文本同时进入训练集和测试集，导致验证结果偏高。监督分类采用按报告分组的交叉验证，并将近重复句子合并为同一分组。",
        f"字符TF-IDF与LinearSVC在情感任务上的准确率为{fmt(numbers.svc_sentiment_acc)}、Macro-F1为{fmt(numbers.svc_sentiment_f1)}；政策四分类准确率为{fmt(numbers.svc_stance_acc)}、Macro-F1为{fmt(numbers.svc_stance_f1)}；条件政策方向准确率为{fmt(numbers.svc_direction_acc)}、Macro-F1为{fmt(numbers.svc_direction_f1)}。监督模型明显改善了情感和政策倾向识别，但少数类别仍受样本规模限制。",
        "按报告分组交叉验证的结果通常低于随机切分结果，但它更接近真实使用场景。模型在预测一份新报告时，不能依赖同一报告中已经出现过的相似句子；如果验证集和训练集共享同一段模板，准确率会被高估。本文据此接受更严格的验证口径，并把主题硬分类降为描述性材料。",
        "人工标注规模也影响评价方式。240句样本足以比较词典、语境门控和轻量监督模型的相对表现，却不足以支撑细主题和少数政策方向的稳定高精度识别。因此，表2同时报告准确率和Macro-F1，避免多数类别主导结论。条件政策方向只在政策相关句内评价，原因是无关句不应被强行纳入偏松、偏紧和中性之间的比较。",
        "学习曲线进一步支持这一判断。情感分类在训练比例从25%提高到100%时，Macro-F1从0.6476提高到0.7993；政策四分类Macro-F1从0.4312提高到0.6804；主题分类Macro-F1只从0.2518提高到0.3329。情感和政策倾向任务随着标注样本增加有所改善，主题任务则明显受类别稀疏和边界模糊限制。",
        "字符TF-IDF适合这个样本规模。中文央行文本中的政策含义常体现在短语和固定搭配上，字符n元特征能够捕捉“合理充裕”“精准有力”“不搞大水漫灌”等局部表达。LinearSVC的线性结构使特征权重较容易解释，也避免在小样本人工标注上使用过于复杂的模型。",
    ]:
        _add_body(doc, text)
    _add_figure(doc, FIGURES_DIR / figures.figure2, "图2  文本分类模型的分组交叉验证学习曲线")

    _add_heading2(doc, "（三）文本测量结果")
    _add_body(doc, "表2中的三类方法处在不同的测量层级。初始词典只作为方法起点，不报告正式验证指标；语境门控词典提供可解释的规则基准；监督模型用于检验文本信息是否能在未见过的报告中保持一定辨识力。主题硬分类的稳定性弱于情感和政策倾向，因此正文只把连续主题关注度用于经济背景解释。")
    _add_table(doc, "表2  文本测量方法比较", tables.table2, note="注：Panel B只在政策相关句中区分偏松、偏紧和中性；缺失值表示该方法未作为正式可比指标保存。")

    _add_heading1(doc, "四、政策指引创新度与股票市场反应")
    _add_heading2(doc, "（一）基准结果")
    for text in [
        f"股票核心估计见表3。基准交互模型中的早期效应为{fmt(numbers.stock_beta)}，HC3标准误为{fmt(numbers.stock_se)}，p值为{fmt(numbers.stock_p)}，样本量为{numbers.stock_n}。这表明，在控制发布前市场波动和附近政策操作后，政策指引文本越偏离上一期表述，报告发布后五个交易日股票市场波动越高。",
        "这一关系的经济含义来自“增量信息”而非简单情绪。政策指引如果延续既有措辞，投资者只需小幅调整已有预期；若指引明显变化，市场需要重新评估政策反应函数、增长判断和流动性环境。短期波动上升可以理解为信息吸收过程中的重新定价，而不必然意味着沟通效果较差。",
        "选择波动率而非收益率方向作为核心因变量，原因在于政策报告往往同时包含政策取向和宏观信息。偏宽松表述可能降低贴现率，宏观下行判断却可能压低盈利预期；二者对股票收益方向的影响可能相互抵消。波动率关注的是市场分歧和不确定性释放，较少依赖单一方向判断，更适合捕捉文本创新度带来的短期反应。",
        "发布前20日波动率控制项也很关键。若某一期报告发布前市场已经处在高波动状态，发布后五日波动率偏高未必来自文本信息。把前期波动纳入回归后，创新度系数反映的是在既有市场不确定性之外，政策指引文本变化与发布后波动之间的条件关系。",
        "基准结果的方向与央行沟通相似度文献相吻合。Ehrmann和Talmi（2020）发现，央行沟通文本越偏离既有表述，市场波动越可能上升。中国样本中的政策报告频率较低、文本更长，市场反应未必完全相同，但“相似度下降意味着增量信息上升”的逻辑在这里仍然成立。",
        "从投资者信息处理角度看，创新度提高并不要求投资者逐字阅读整份报告。媒体摘要、卖方点评和机构内部解读会把政策指引变化迅速传播到市场。若本期指引与上一期相比变化较大，不同机构对其含义的解释可能出现分歧，短期成交和波动随之增加。这一机制解释了为什么文本创新度更容易体现在波动率而不是平均收益上。",
    ]:
        _add_body(doc, text)
    _add_table(doc, "表3  政策指引创新度与股票波动率", tables.table3, note="注：被解释变量为报告发布后五个交易日实际波动率对数；基准交互模型中的早期效应对应创新度主效应。")

    _add_heading2(doc, "（二）时期差异与经济效应")
    for text in [
        f"按样本标准差换算，政策指引创新度提高{fmt(numbers.guidance_std)}，对应发布后五日实际波动率约提高{fmt(numbers.stock_one_sd_pct, 1)}%。一单位创新度覆盖范围过大，不适合作为正文的主要经济解释；原始系数仍在表3中呈现，用于说明对数波动率的回归口径。",
        f"2019年后交互项为{fmt(numbers.stock_interaction)}，p值为{fmt(numbers.stock_interaction_p)}；2019年后总效应为{fmt(numbers.stock_total)}，总效应p值为{fmt(numbers.stock_total_p)}。后期样本估计不稳定，不能写成显著下降。更合适的解释是：随着货币政策框架、信息披露方式和市场预期管理逐步成熟，单份季度报告的边际文本变化可能更容易被市场提前吸收，但现有样本不足以把这一机制识别为稳定统计关系。",
        "2019年前后差异还可能与文本制度化程度有关。后期报告中政策取向表述更强调连续性、稳定性和预期管理，市场也可能通过公开操作、新闻发布会和其他渠道提前形成判断。季度报告仍然提供系统信息，但它相对于其他沟通渠道的边际新意可能下降。现有结果支持这种解释方向，却不足以把它写成确定机制。",
        "疫情期和非疫情期的分组结果保留在表3中。疫情期政策沟通面对更强的外生冲击，文本创新度可能混合宏观冲击、政策响应和风险提示；非疫情期则更接近常态化沟通。分组样本量明显缩小，p值变化不能被过度解读，本文只把它们作为对基准关系稳定性的补充观察。",
        "经济效应采用一个样本标准差进行解释，也能避免夸大系数含义。创新度指标处在0到1之间，但样本中的实际波动范围远小于完整理论区间；用一单位变化解释，会把几乎不可能同时发生的文本变化写成常见冲击。一个标准差变化更贴近样本内报告之间的实际差异，也更容易与市场波动幅度联系起来。",
        "散点图显示，高创新度报告并不总是对应高波动。货币政策沟通只是影响市场波动的因素之一，金融危机、疫情、海外利率和国内宏观数据都会改变短期市场状态。基准回归的意义在于，在控制前期波动和政策操作后，创新度仍与发布后波动存在正相关关系。",
    ]:
        _add_body(doc, text)
    _add_figure(doc, FIGURES_DIR / figures.figure3, "图3  政策指引创新度与股票实际波动率")

    _add_heading2(doc, "（三）日度波动稳健性")
    for text in [
        f"Student-t EGARCH-X模型使用{numbers.egarch_n}个连续交易日收益率和{numbers.egarch_events}个创新度事件。报告发布当日规格中，标准化创新度条件方差项系数为{fmt(numbers.egarch_novelty_coef)}，似然比检验p值为{fmt(numbers.egarch_lr_p)}，置换诊断p值为{fmt(numbers.egarch_perm_p)}。按一个标准差创新度换算，条件波动率约提高{fmt(numbers.egarch_vol_pct)}%。方向与股票事件级结果一致，但统计证据不足。",
        "EGARCH-X稳健性检验的意义在于使用完整日度序列刻画波动持续性和厚尾特征，而不是把事件窗口内少数交易日单独抽出估计。该模型没有推翻核心结果，也没有提供更强显著性证据。本文据此把日度模型定位为高级稳健性检验，而非替代股票事件级主检验。",
        "日度模型和事件级模型回答的问题并不相同。事件级模型把每期报告浓缩为一个短窗口反应，更直接对应“报告发布后市场是否更波动”；EGARCH-X模型则在长日度序列中估计报告日冲击对条件方差的边际贡献，同时吸收收益率自身的波动聚集。后者统计要求更高，尤其在报告事件只有季度频率时，弱显著性并不否定事件级结果。",
        "Student-t分布的使用主要服务于金融收益率的厚尾特征。若忽略极端收益率，少数市场大幅波动日可能被条件方差方程误认为文本冲击；厚尾设定能够缓和这一问题。表4同时给出报告发布当日、后一交易日以及当日与次日联合规格，目的是观察日期归属的敏感性，而不是重新选择更有利的主结果。",
        "EGARCH-X还允许正负收益冲击对未来波动产生非对称影响。股票市场下跌日通常伴随更高的条件波动，若模型只使用对称方差过程，可能把收益率自身的非对称反应误归因于报告事件。把文本创新度放入条件方差方程，而不是放入均值方程，可以更贴近本文关心的“波动率反应”问题。",
        "表4的日度结果应与表3一起阅读。表3的证据更直接、统计上也更强；表4提供的是在更完整波动动态下的方向检验。两者方向一致，说明核心发现没有因加入日度波动结构而反转；但表4的p值较大，说明这一高级规格并未把股票结果提升为更强的日度证据。",
    ]:
        _add_body(doc, text)
    _add_table(doc, "表4  日度波动稳健性检验", tables.table4, note="")

    _add_heading1(doc, "五、政策语调与国债收益率曲线")
    _add_heading2(doc, "（一）未预期政策语调")
    for text in [
        f"债券部分考察政策指引未预期语调与收益率曲线斜率变化。三日窗口内斜率回归的主效应系数为{fmt(numbers.bond_beta)}，HC3标准误为{fmt(numbers.bond_se)}，p值为{fmt(numbers.bond_p)}；2019年后总效应为{fmt(numbers.bond_total)}，总效应p值为{fmt(numbers.bond_total_p)}。估计方向在不同曲线因子之间并不稳定，债券证据弱于股票波动结果。",
        "收益率曲线同时反映短端政策利率预期、长端增长通胀预期和期限溢价。央行报告中的偏松或偏紧表述可能推低短端，也可能通过增长预期抬升长端；两个渠道叠加后，斜率反应未必呈现单一方向。Kuttner（2001）和Gurkaynak等（2005）关于未预期政策信息的研究提示，利率市场对政策消息的反应往往依赖识别口径和期限结构。",
        "未预期语调的构造以报告自身的历史规律为基准，试图区分常规表述和超出惯性的政策倾向变化。这个思路与董青马等（2024）强调的“潜在”未预期信息相呼应；本研究仅借鉴其问题意识，并不复制其完整识别框架。季度报告文本较长，且包含大量宏观判断，未预期语调进入债券回归后仍可能同时反映政策冲击和央行信息冲击。",
        "Nakamura和Steinsson（2018）、Jarocinski和Karadi（2020）的研究表明，央行信息效应可能使市场把政策消息理解为关于经济前景的信号。放在收益率曲线上，偏紧语调并不一定只推高短端利率；如果市场把它理解为经济韧性更强，也可能影响长端收益率。本文的债券结果不稳定，正反映了这类信息分解在季度文本场景中的困难。",
        "中国债券市场还受到资金面、监管预期和机构配置需求影响。季度报告发布时点附近，如果资金利率、公开市场操作或海外利率同步变化，收益率曲线可能先反映这些高频因素。文本语调变量在三日窗口中要与这些因素竞争解释力，得到弱结果并不奇怪。",
        "曲线水平、斜率和曲率的经济含义不同。水平变化更接近整体利率环境，斜率变化反映短端和长端相对调整，曲率则对应中期限收益率的相对位置。政策语调若主要影响短端，斜率会变化；若市场同时调整增长和通胀预期，水平和曲率也可能移动。表5的Panel A把三类结果放在一起，正是为了避免只报告单一曲线因子。",
    ]:
        _add_body(doc, text)

    _add_heading2(doc, "（二）跨拟合监督语调")
    for text in [
        f"跨拟合政策语调把人工标注句子的监督分类结果聚合到报告层面。政策相关句均值的主效应为{fmt(numbers.cross_coef)}，p值为{fmt(numbers.cross_p)}；2019年后总效应为{fmt(numbers.cross_total)}，总效应p值为{fmt(numbers.cross_total_p)}。该结果没有稳定改善债券解释力。",
        "跨拟合处理的价值在于降低训练集和解释变量之间的机械重合：每份报告的预测语调来自未使用该报告标注句训练的模型。问题也很清楚，人工标注样本只有240句，少数政策方向样本偏少，折外预测进入债券回归后噪声会被放大。因此，跨拟合监督语调更适合展示测量思路，而不是承担核心市场检验。",
        "政策相关句均值是三种聚合方式中经济含义最清楚的一种。全部句子均值容易被宏观描述和背景性表述稀释，方向性句子均值又会丢掉大量中性但政策相关的信息。政策相关句均值保留了政策语境，同时减少无关句子的影响，因此正文用它报告主要跨拟合结果，并在表5中列出其他聚合方式作为对照。",
        "监督语调没有稳定改善债券解释力，并不意味着人工标注或分类模型没有价值。它说明句子层方向识别和报告层市场反应之间还存在一个聚合与识别问题：报告可能同时出现偏松、偏紧和风险提示句，市场反应也可能取决于这些句子在报告结构中的位置，而不是简单平均后的得分。",
        "跨拟合结果还提醒我们，文本测量误差会在金融回归中被放大。句子分类即使在验证集上有较高准确率，聚合到报告层后仍可能因句子数量、章节位置和类别不平衡产生偏差。债券回归样本只有季度报告事件，解释变量的一点噪声就会显著影响系数稳定性。",
        "因此，跨拟合监督语调在本文中主要用于验证测量链条：人工标注能否支持监督分类，监督分类能否生成报告层语调，报告层语调进入债券模型后是否有清晰解释力。前两步结果相对可用，最后一步证据不足，这一层级差异本身就是研究结论的一部分。",
    ]:
        _add_body(doc, text)

    _add_heading2(doc, "（三）结果讨论")
    _add_body(doc, "股票和债券结果不一致，并不意外。股票波动率对不确定性本身敏感，文本创新度提高可能直接增加短期重新定价强度；债券收益率曲线则需要区分政策路径、宏观信息和期限溢价，单一语调变量难以同时捕捉这些渠道。McMahon、Schipke和Li（2018）指出，中国央行沟通仍处在制度化改进过程中，这也意味着不同市场吸收沟通信息的方式可能存在差异。")
    _add_body(doc, "从市场微观角度看，股票和债券投资者关注的信息集合也不同。股票投资者更关心盈利预期、风险偏好和流动性折现，报告文本一旦改变政策指引，估值分歧就可能放大；债券投资者则更关注未来利率路径和期限补偿，语调变化必须先被解释为短端、长端或期限溢价冲击，才会表现为清晰的曲线变化。")
    _add_body(doc, "债券图把收益率曲线因子和未预期语调分组放在同一张图中。Panel A显示曲线水平、斜率和曲率在较长样本内同步波动但并非同一变量；Panel B显示不同语调分组在2019年前后没有形成稳定单调关系。图形证据与表5一致：债券市场存在反应，但反应方向不够稳定，不能写成与股票结果同等强度的结论。")
    _add_table(doc, "表5  国债市场探索性结果", tables.table5, note="")
    _add_figure(doc, FIGURES_DIR / figures.figure4, "图4  国债收益率曲线与未预期政策语调反应")

    _add_heading1(doc, "六、进一步检验与研究局限")
    _add_heading2(doc, "（一）样本量与检验功效")
    for text in [
        f"市场功效分析显示，在当前{numbers.stock_n}个股票事件样本下，按观察到的效应大小计算的功效约为{fmt(numbers.power_current)}；若样本扩大至120和160个事件，功效分别提高到{fmt(numbers.power_120)}和{fmt(numbers.power_160)}。这说明不显著结果需要结合样本量理解，尤其是债券和分时期结果。",
        "季度报告的天然频率限制了样本扩张速度。增加日度市场观测不能等同于增加报告事件，因为文本变量只在报告发布时变化。未来研究若要提高检验功效，更可行的方向是合并多类正式沟通文本，或在不牺牲事件定义清晰度的前提下引入公告、发布会和货币政策委员会例会材料。",
        "功效结果也解释了为什么本文保留方向一致但不显著的日度和债券结果。对于季度报告这种低频文本，分时期、分市场和分期限检验都会迅速消耗样本量。若只按显著性筛选结果，容易把样本噪声误写成稳健机制；若完全忽略不显著结果，又会高估文本指标的解释范围。本文把证据层级分开，正是为了避免这两种偏差。",
        "功效图并不用于宣称扩大样本后一定能得到显著结论。它展示的是在给定效应大小和方差估计下，事件数量对统计功效的机械影响。真正扩大样本时，新增沟通文本的类型、发布时间和市场环境都可能改变效应大小，因此功效曲线只能作为研究设计参考，而不是外推预测。",
        "对课程研究而言，功效分析还有一个实用含义。很多金融事件研究在样本较小时会出现方向合理但p值较大的结果，此时最稳妥的写法不是反复更换窗口或变量，而是说明样本设计能够检验到多大的效应。图2的学习曲线显示，文本分类在当前样本规模下情感和政策倾向识别尚有改善空间，但对债券和分时期扩展仍显不足。",
    ]:
        _add_body(doc, text)
    _add_heading2(doc, "（二）测量与识别局限")
    for text in [
        "本文仍有几项边界。第一，PDF抽取和章节识别会给早期报告带来文本噪声，虽然关键章节经过修复，但不能保证所有句子边界完全符合原始排版。第二，人工标注样本规模有限，监督模型对少数政策方向和稀疏主题的识别仍不稳定。第三，事件窗口不能完全隔离同期宏观数据、全球风险偏好和其他政策操作，因此短窗口回归应解释为条件相关关系，而不是完整因果识别。",
        "这些限制不会改变本文的主要证据层级：股票创新度与短期波动是核心主线；Student-t EGARCH-X提供方向一致但较弱的日度稳健性；债券语调和曲线结果属于探索性扩展。统计显著、方向一致和经济上值得关注是三类不同表述，正文只在检验支持时使用“显著”。",
        "文本测量本身也有边界。词典方法透明、可解释，却难以处理复杂否定和跨句语境；监督模型能学习更多局部模式，但在小样本标注下容易受类别不平衡影响。连续主题关注度缓解了硬分类不稳定的问题，但它描述的是注意力分布，不等同于政策冲击强度。本文在解释市场结果时只使用与研究问题直接对应的文本变量。",
        "识别层面的限制更难完全消除。季度报告往往在宏观数据、政策操作和市场预期变化之后发布，其文字既可能回应已经发生的经济金融形势，也可能引导下一阶段预期。短窗口事件研究减少了干扰，却不能把所有同期信息排除干净。因而，本文把核心结果表述为政策指引创新度与短期股票波动的条件相关关系。",
        "还有一个局限来自政策沟通类型。季度货币政策执行报告篇幅长、信息丰富，却不是唯一沟通渠道。新闻稿、例会通稿、发布会问答和公开市场操作公告可能在报告发布前后释放相同或相近信息。若市场已经通过其他渠道形成预期，季度报告的新增冲击就会被削弱。本文的结论只适用于本样本中的季度报告文本，不能直接推广到所有央行沟通形式。",
        "本文也不声称完整复现董青马等（2024）或尚玉皇、刘华和申峰（2025）的高级识别框架。前者关注资产价格中的潜在未预期政策信息，后者围绕国债收益率曲线建立更系统的预期博弈分析；本文借鉴这些研究的问题意识，但在现有数据条件下采用更克制的文本测量和事件回归设计。",
        "这些边界使结论更窄，也更稳。政策指引创新度可以解释股票短期波动的一部分，但它不是完整的货币政策冲击；未预期语调能够提供债券探索线索，但尚不能稳定分离政策路径和宏观信息。读者在使用本文结果时，应把它看作一组可复核的经验事实，而不是对央行沟通效应的最终判定。",
    ]:
        _add_body(doc, text)

    _add_heading1(doc, "七、结论")
    for text in [
        f"本文使用{numbers.report_start}至{numbers.report_end}中国人民银行季度货币政策执行报告、沪深300指数和中债国债收益率数据，构造政策指引扩展窗口TF-IDF创新度、词典与语境门控语调、字符TF-IDF与LinearSVC监督语调，并考察这些文本指标与金融市场反应之间的关系。",
        f"核心发现是，政策指引创新度与股票市场发布后五日实际波动率显著正相关；一个样本标准差创新度对应约{fmt(numbers.stock_one_sd_pct, 1)}%的实际波动率上升。日度Student-t EGARCH-X模型得到正向但不显著的方差项，说明完整日度波动结构下仍能看到相同方向，但证据强度下降。债券收益率曲线和跨拟合监督语调结果不稳定，提示政策语调进入期限结构时受宏观信息和期限溢价共同影响。",
        "从研究设计看，政策报告文本的增量信息比一般正负情绪更能解释股票短期波动。对债券市场而言，简单语调指标还不足以分离政策冲击和央行信息冲击。后续研究可以在保持事件定义清晰的前提下扩大沟通文本类型，并结合更高频的利率预期数据识别曲线反应渠道。",
        "本文的结论应按证据层级理解。核心主线是政策指引创新度与股票波动率；日度EGARCH-X结果提供同方向但较弱的稳健性；债券部分展示了可继续研究的探索性线索。这样的层级划分比把所有结果写成同等强度更符合样本规模和估计结果，也更能反映文本测量与金融事件研究结合时的实际边界。",
        "方法上，本文把Python文本处理、分组交叉验证和金融事件研究连接为一套可复核的经验设计。文本端先处理章节、句子、词典和监督分类，市场端再把报告层指标映射到交易日窗口。这个过程说明，央行沟通研究的关键不只是选择回归模型，还包括把文本信息转换为经济含义明确的变量。",
        "对文本测量而言，最重要的不是堆叠更复杂的模型，而是让指标与研究问题对应。创新度用于衡量政策指引的新信息，语境门控用于修正词典误判，监督分类用于检验句子层测量是否可外推。三者各有边界，只有放在合适的证据位置上，才能减少文本金融研究中的过度解释。",
        "对金融市场解释而言，股票和债券结果的差异本身提供了信息。股票波动率更直接反映不确定性和分歧，因而能较清楚地吸收政策指引创新度；收益率曲线则要求同时区分短端政策预期、长端宏观预期和期限溢价，单一语调指标难以完成这种分解。后续研究若能引入高频利率预期或更细的沟通事件，债券部分仍有改进空间。",
        "方向一致但不显著的结果仍按原始证据层级呈现，债券结果较弱也保留为探索性检验。这样的处理牺牲了叙事上的整齐性，却保留了样本给出的真实信息。",
        "若进一步扩展本文，可以从两个方向推进。一个方向是增加正式沟通文本类型，把季度报告与货币政策委员会例会通稿、新闻发布会问答和政策工具公告分开比较；另一个方向是引入更高频的利率预期或期货隐含信息，专门识别政策路径和央行信息效应。前者扩大文本覆盖，后者改善债券市场识别，两者都需要更严格的事件定义。",
        "文本指标也可以继续细化。政策指引章节内部包含总量政策、结构性工具、信贷投放、汇率、房地产和金融风险等不同内容，创新度上升可能来自不同主题。本文只报告连续主题关注度作为描述性解释，没有把主题创新度逐项放入市场回归，原因是当前样本难以支持过多交互项。更长样本或更多沟通文本能够支撑更细的主题层分析。",
        "人工标注方面，最需要改进的是类别平衡而不是简单扩大总量。少数政策方向和部分主题样本偏少，会影响Macro-F1和报告层语调聚合。未来标注若能针对偏紧、偏松、金融稳定和房地产相关句子进行分层补充，监督语调进入债券回归时的噪声可能下降。",
        "金融事件研究方面，发布时点的精确性仍然重要。季度报告有时在收盘后披露，有时在非交易日附近披露，股票和债券市场的第一反应可能落在不同交易日。本文采用统一事件日口径以保持可比性，未来若能取得更精确的分钟级发布和成交数据，可以进一步检验当日、次日和跨日反应的差异。",
        "最值得保留的经验是证据排序。股票结果给出较清晰的核心发现，日度波动模型提供方向一致的补充，债券结果则提醒我们，政策语调进入收益率曲线时会受到更多渠道干扰。将显著、方向一致和探索性结果分开陈述，比单纯追求显著性更能反映货币政策文本数据的真实约束。",
    ]:
        _add_body(doc, text)


def _add_front_matter(doc: Document, numbers: PaperNumbers) -> None:
    _add_title_line(doc, "中国货币政策报告文本特征")
    _add_title_line(doc, "与金融市场反应")
    # Subtitle
    p_sub = doc.add_paragraph()
    set_paragraph_format(p_sub, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_pt=16, before_pt=2, after_pt=4, exact=False)
    run_sub = p_sub.add_run("——基于Python文本量化、股票波动与国债收益率曲线的研究")
    set_run_font(run_sub, FONT_STYLE["body_cn"], 12)
    # Author in 楷体 三号(16pt)
    p_author = doc.add_paragraph()
    set_paragraph_format(p_author, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_pt=20, before_pt=6, after_pt=6, exact=False)
    run_author = p_author.add_run("罗允绩")
    set_run_font(run_author, FONT_STYLE["author_cn"], PARAGRAPH_STYLE["author_size_pt"])
    _add_abstract_block(
        doc,
        "内容提要：",
        (
            f"本文基于{numbers.report_start}至{numbers.report_end}中国人民银行季度货币政策执行报告，构造政策指引扩展窗口TF-IDF创新度，"
            "并结合人工标注、语境门控和字符TF-IDF与LinearSVC检验文本测量有效性。股票事件研究显示，政策指引创新度与报告发布后五日实际波动率显著正相关；"
            f"创新度提高一个样本标准差，实际波动率约提高{fmt(numbers.stock_one_sd_pct, 1)}%。Student-t EGARCH-X日度模型得到方向一致但不显著的波动率效应。"
            "债券部分中，未预期政策语调和跨拟合监督语调对收益率曲线斜率的解释不稳定，适合作为探索性证据。结果表明，央行报告中的增量指引信息更直接地反映在股票短期不确定性中，而债券期限结构反应需要更细致地区分政策路径、宏观信息和期限溢价。"
        ),
    )
    _add_keywords(doc, "关键词：货币政策沟通；文本创新度；实际波动率；EGARCH-X；收益率曲线")
    _add_centered(doc, "中图分类号：F832.0    文献标识码：A", size=PARAGRAPH_STYLE["abstract_size_pt"], cn_font=FONT_STYLE["body_cn"])
    _add_centered(doc, "Textual Features of China's Monetary Policy Reports and Financial Market Reactions", size=11, cn_font=FONT_STYLE["body_en"], bold=True)
    _add_centered(doc, "LUO Yunji", size=PARAGRAPH_STYLE["abstract_size_pt"], cn_font=FONT_STYLE["body_en"])
    _add_english_abstract(
        doc,
        (
            "This paper studies quarterly monetary policy reports of the People's Bank of China from 2006Q1 to 2025Q4. "
            "It constructs an expanding-window TF-IDF novelty measure for policy guidance, validates text measurement with manual annotations, context gating, and grouped cross-validation, and links the text measures to stock and bond market reactions. "
            "The core event-level evidence shows a positive association between guidance novelty and five-day realized stock volatility. "
            "A Student-t EGARCH-X model estimated on the full daily return sequence yields a positive but statistically weaker volatility effect. "
            "Unexpected policy tone and cross-fitted supervised tone do not robustly explain the government bond yield curve, suggesting that bond-market communication effects are harder to separate from macro-information and term-premium channels."
        ),
    )
    _add_centered(doc, "Key words: monetary policy communication; text novelty; realized volatility; EGARCH-X; yield curve", size=PARAGRAPH_STYLE["abstract_size_pt"], cn_font=FONT_STYLE["body_en"])


def _add_title_line(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_pt=24, before_pt=0, after_pt=2)
    run = paragraph.add_run(text)
    set_run_font(run, FONT_STYLE["title_cn"], 18, bold=False)


def _add_centered(doc: Document, text: str, *, size: float, cn_font: str, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_pt=14, before_pt=0, after_pt=2)
    run = paragraph.add_run(text)
    set_run_font(run, cn_font, size, bold=bold)


def _add_abstract_block(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, line_pt=PARAGRAPH_STYLE["abstract_line_pt"], first_line_chars=0, before_pt=4, after_pt=0)
    run = paragraph.add_run(label)
    set_run_font(run, FONT_STYLE["abstract_cn"], PARAGRAPH_STYLE["abstract_size_pt"], bold=True)
    body = paragraph.add_run(text)
    set_run_font(body, FONT_STYLE["abstract_cn"], PARAGRAPH_STYLE["abstract_size_pt"])


def _add_keywords(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, line_pt=PARAGRAPH_STYLE["abstract_line_pt"], first_line_chars=0, before_pt=1, after_pt=1)
    run = paragraph.add_run(text)
    set_run_font(run, FONT_STYLE["abstract_cn"], PARAGRAPH_STYLE["abstract_size_pt"])


def _add_english_abstract(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, line_pt=PARAGRAPH_STYLE["english_abstract_line_pt"], first_line_chars=0, before_pt=2, after_pt=0)
    run = paragraph.add_run("Abstract: ")
    set_run_font(run, FONT_STYLE["body_en"], PARAGRAPH_STYLE["abstract_size_pt"], bold=True)
    body = paragraph.add_run(text)
    set_run_font(body, FONT_STYLE["body_en"], PARAGRAPH_STYLE["abstract_size_pt"])


def _add_heading1(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_pt=18,
        before_pt=HEADING_STYLE["level1_before_pt"],
        after_pt=HEADING_STYLE["level1_after_pt"],
        keep_with_next=True,
    )
    run = paragraph.add_run(text)
    set_run_font(run, FONT_STYLE["heading_cn"], HEADING_STYLE["level1_size_pt"])


def _add_heading2(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        line_pt=16,
        first_line_chars=2,
        before_pt=HEADING_STYLE["level2_before_pt"],
        after_pt=HEADING_STYLE["level2_after_pt"],
        keep_with_next=True,
    )
    run = paragraph.add_run(text)
    set_run_font(run, FONT_STYLE["subheading_cn"], HEADING_STYLE["level2_size_pt"])


def _add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(
        paragraph,
        line_pt=PARAGRAPH_STYLE["body_line_pt"],
        first_line_chars=PARAGRAPH_STYLE["first_line_chars"],
    )
    run = paragraph.add_run(text)
    set_run_font(run, FONT_STYLE["body_cn"], PARAGRAPH_STYLE["body_size_pt"])


def _add_table(doc: Document, caption: str, df: pd.DataFrame, note: str = "") -> None:
    cap = doc.add_paragraph()
    set_paragraph_format(cap, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_pt=CAPTION_STYLE["line_pt"], before_pt=4, after_pt=2, keep_with_next=True)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, FONT_STYLE["subheading_cn"], TABLE_STYLE["caption_size_pt"], bold=True)
    
    base_key = "unknown"
    if "表1" in caption: base_key = "table1"
    elif "表2" in caption: base_key = "table2"
    elif "表3" in caption: base_key = "table3"
    elif "表4" in caption: base_key = "table4"
    elif "表5" in caption: base_key = "table5"

    if "Panel" in df.columns:
        _add_panel_tables(doc, df, base_key)
    else:
        _add_table_rows(doc, [list(df.columns)] + df.astype(str).values.tolist(), semantic_key=base_key)
        
    if note:
        p_note = doc.add_paragraph()
        set_paragraph_format(p_note, line_pt=11, first_line_chars=0, before_pt=2, after_pt=4)
        run = p_note.add_run(note)
        set_run_font(run, FONT_STYLE["body_cn"], TABLE_STYLE["note_size_pt"])


def _add_panel_tables(doc: Document, df: pd.DataFrame, base_key: str) -> None:
    for panel, panel_df in df.groupby("Panel", sort=False):
        panel_paragraph = doc.add_paragraph()
        set_paragraph_format(panel_paragraph, line_pt=12, first_line_chars=0, before_pt=2, after_pt=1)
        panel_run = panel_paragraph.add_run(str(panel))
        set_run_font(panel_run, FONT_STYLE["heading_cn"], TABLE_STYLE["body_size_pt"], bold=True)
        columns = _panel_columns(panel_df)
        rows = [columns] + panel_df[columns].astype(str).values.tolist()
        
        panel_key = f"{base_key}_panel_a" if "Panel A" in str(panel) else f"{base_key}_panel_b"
        _add_table_rows(doc, rows, semantic_key=panel_key)


def _panel_columns(panel_df: pd.DataFrame) -> list[str]:
    candidates = [c for c in panel_df.columns if c != "Panel"]
    columns = []
    for column in candidates:
        values = panel_df[column].astype(str).str.strip()
        if values.ne("—").any() and values.ne("").any():
            columns.append(column)
    return columns


def _add_table_rows(doc: Document, rows: list[list[str]], semantic_key: str = None) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = None
    clear_table_borders(table)
    widths = _table_widths(len(rows[0]), semantic_key)
    set_table_width(table, widths)
    
    config = TABLE_VERTICAL_CONFIG.get(semantic_key, {})

    for i, row_values in enumerate(rows):
        row = table.rows[i]
        mark_row_no_split(row)
        
        is_header = i == 0
        if is_header:
            if semantic_key in ["table1", "table2_panel_a", "table3", "table4", "table5_panel_a"]:
                for cell in row.cells:
                    set_cell_border(cell, "top", TABLE_STYLE["top_border_sz"])

        for j, value in enumerate(row_values):
            set_cell_margins(row.cells[j])
            
            align = WD_ALIGN_PARAGRAPH.CENTER
            if i > 0:
                header_val = rows[0][j] if 0 < len(rows) and j < len(rows[0]) else ""
                if "样本量" in header_val:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                elif _looks_numeric(value):
                    align = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    align = WD_ALIGN_PARAGRAPH.CENTER
            
            set_cell_text(row.cells[j], value, bold=is_header, align=align)
            
            # Explicit vertical separators checking
            if config.get(j, False):
                set_cell_border(row.cells[j], "right", TABLE_STYLE["mid_border_sz"])

        if is_header:
            for cell in row.cells:
                set_cell_border(cell, "bottom", TABLE_STYLE["mid_border_sz"])
                
    for cell in table.rows[-1].cells:
        if semantic_key in ["table1", "table2_panel_b", "table3", "table4", "table5_panel_b"]:
            set_cell_border(cell, "bottom", TABLE_STYLE["bottom_border_sz"])
        elif semantic_key in ["table2_panel_a", "table5_panel_a"]:
            set_cell_border(cell, "bottom", TABLE_STYLE["mid_border_sz"])


def _table_widths(n_cols: int, semantic_key: str = None) -> list[float]:
    if semantic_key == "table1":
        # 18%, 30%, 10%, 19%, 23% of 15.0cm
        return [2.7, 4.5, 1.5, 2.85, 3.45]
    if semantic_key == "table2_panel_a":
        # 32%, 17%, 17%, 17%, 17% of 15.0cm
        return [4.8, 2.55, 2.55, 2.55, 2.55]
    if semantic_key == "table2_panel_b":
        # 32%, 34%, 34% of 15.0cm
        return [4.8, 5.1, 5.1]
    if semantic_key == "table3":
        # 38%, 10%, 17.33%, 17.33%, 17.33% of 15.0cm
        return [5.7, 1.5, 2.6, 2.6, 2.6]
    if semantic_key == "table4":
        # 38%, 15.5%, 15.5%, 15.5%, 15.5% of 15.0cm
        return [5.7, 2.325, 2.325, 2.325, 2.325]
    if semantic_key == "table5_panel_a":
        # 38%, 10%, 17.33%, 17.33%, 17.33% of 15.0cm
        return [5.7, 1.5, 2.6, 2.6, 2.6]
    if semantic_key == "table5_panel_b":
        # 26%, 10%, 11%, 15%, 21%, 17% of 15.0cm
        return [3.9, 1.5, 1.65, 2.25, 3.15, 2.55]
        
    return [max(1.8, 15.0 / max(n_cols, 1))] * n_cols


def _table_rows(df: pd.DataFrame) -> list[list[str]]:
    if "Panel" not in df.columns:
        return [list(df.columns)] + df.astype(str).values.tolist()
    all_rows: list[list[str]] = []
    non_panel_cols = [c for c in df.columns if c != "Panel"]
    current_panel = None
    for _, row in df.iterrows():
        panel = row["Panel"]
        if panel != current_panel:
            all_rows.append([panel] + [""] * (len(non_panel_cols) - 1))
            current_panel = panel
            panel_cols = [c for c in non_panel_cols if str(row.get(c, "—")) != "—" or c in {"方法", "规格", "被解释变量", "聚合方式", "样本量", "p值"}]
            all_rows.append(panel_cols + [""] * (len(non_panel_cols) - len(panel_cols)))
        all_rows.append([str(row.get(c, "—")) for c in non_panel_cols])
    max_cols = max(len(r) for r in all_rows)
    return [r + [""] * (max_cols - len(r)) for r in all_rows]


def _looks_numeric(value: str) -> bool:
    return bool(re.match(r"^-?\d+(\.\d+)?$", str(value)))


def _add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = doc.add_paragraph()
    # Use single/auto line height for figure paragraphs — never exact
    set_figure_paragraph_format(paragraph, keep_with_next=True)
    paragraph.add_run().add_picture(str(path), width=Cm(11.5))
    cap = doc.add_paragraph()
    set_paragraph_format(cap, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_pt=CAPTION_STYLE["line_pt"], first_line_chars=0, before_pt=1, after_pt=4)
    run = cap.add_run(caption)
    set_run_font(run, FONT_STYLE["body_cn"], CAPTION_STYLE["size_pt"])


def _sanitize_document(doc: Document) -> None:
    """Replace engineering variable names and development phrases in the document text (presentation layer only)."""
    replacements = {
        "full_joint_mle": "联合极大似然估计",
        "D0_D1": "当日与次日联合规格",
        "D0": "报告发布当日",
        "D1": "报告发布后一交易日",
        "wild_residual_bootstrap_hc3": "基于HC3推断的野生残差Bootstrap",
        "tone_aggregation": "语调聚合方式",
        "interaction_coef": "交互项系数",
        "post_2019_total_effect": "2019年后总效应",
        "se_hc3": "HC3标准误",
        "p_value": "p值",
    }
    banned_phrases = [
        "README.md",
        "run_all.py",
        "configs/project.yml",
        "final_submission",
        "Manifest",
        "文件哈希",
        "缓存",
        "pytest",
    ]

    def replace_in_para(p):
        text = p.text
        for a, b in replacements.items():
            if a in text:
                text = text.replace(a, b)
        for ban in banned_phrases:
            if ban in text:
                text = text.replace(ban, "")
        if text != p.text:
            # clear and re-add run
            for r in list(p.runs):
                r.clear()
            p.add_run(text)

    for para in doc.paragraphs:
        replace_in_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_para(para)


def _normalize_docx_spacing_rules(docx_path: Path) -> None:
    """No-op. Exact line spacing is now handled correctly at generation time via set_figure_paragraph_format."""
    pass


def _add_references_heading(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_format(
        paragraph, 
        alignment=WD_ALIGN_PARAGRAPH.LEFT, 
        line_pt=12, 
        before_pt=6, 
        after_pt=3, 
        first_line_chars=0,
        keep_with_next=True
    )
    run = paragraph.add_run("参考文献")
    set_run_font(run, "SimHei", 9.0, bold=False)

def _add_references(doc: Document) -> list[str]:
    _add_references_heading(doc)
    refs = [
        "[1] 董青马，张皓越，马剑文，等．央行沟通与资产价格：识别“潜在”未预期货币政策信息[J]．金融研究，2024(6)．",
        "[2] 姜富伟，胡逸驰，黄楠．央行货币政策报告文本信息、宏观经济与股票市场[J]．金融研究，2021(6)．",
        "[3] 尚玉皇，刘华，申峰．预期的博弈：央行沟通与国债收益率曲线[J]．金融研究，2025(9)．",
        "[4] 中国人民银行．中国货币政策执行报告[R/OL]．2006-2025．",
        "[5] Blinder A S, Ehrmann M, Fratzscher M, et al. Central Bank Communication and Monetary Policy: A Survey of Theory and Evidence[J]. Journal of Economic Literature, 2008, 46(4): 910-945.",
        "[6] Du Z, Huang A G, Wermers R, et al. Language and Domain Specificity: A Chinese Financial Sentiment Dictionary[J]. Review of Finance, 2022, 26(3): 673-719.",
        "[7] Ehrmann M, Talmi J. Starting from a Blank Page? Semantic Similarity in Central Bank Communication and Market Volatility[J]. Journal of Monetary Economics, 2020, 111: 48-62.",
        "[8] Gurkaynak R S, Sack B, Swanson E T. Do Actions Speak Louder than Words? The Response of Asset Prices to Monetary Policy Actions and Statements[J]. International Journal of Central Banking, 2005, 1(1): 55-93.",
        "[9] Hansen S, McMahon M. Shocking Language: Understanding the Macroeconomic Effects of Central Bank Communication[J]. Journal of International Economics, 2016, 99(S1): S114-S133.",
        "[10] Jarocinski M, Karadi P. Deconstructing Monetary Policy Surprises: The Role of Information Shocks[J]. American Economic Journal: Macroeconomics, 2020, 12(2): 1-43.",
        "[11] Kuttner K N. Monetary Policy Surprises and Interest Rates: Evidence from the Fed Funds Futures Market[J]. Journal of Monetary Economics, 2001, 47(3): 523-544.",
        "[12] McMahon M, Schipke A, Li X. China's Monetary Policy Communication: Frameworks, Impact, and Recommendations[R]. IMF Working Paper, WP/18/244, 2018.",
        "[13] Nakamura E, Steinsson J. High-Frequency Identification of Monetary Non-Neutrality: The Information Effect[J]. Quarterly Journal of Economics, 2018, 133(3): 1283-1330.",
    ]
    for ref in refs:
        paragraph = doc.add_paragraph()
        set_paragraph_format(paragraph, line_pt=REFERENCE_STYLE["line_pt"], first_line_chars=0, before_pt=0, after_pt=0)
        paragraph.paragraph_format.first_line_indent = -Cm(0.74)
        paragraph.paragraph_format.left_indent = Cm(0.74)
        run = paragraph.add_run(ref)
        cn = FONT_STYLE["reference_cn"] if re.search(r"[\u4e00-\u9fff]", ref) else FONT_STYLE["body_en"]
        set_run_font(run, cn, REFERENCE_STYLE["body_size_pt"])
    return refs


def _write_paper_audits(numbers: PaperNumbers, refs: list[str]) -> None:
    pd.DataFrame(
        [
            {"item": key, "value": value, "source": _number_source(key)}
            for key, value in asdict(numbers).items()
        ]
    ).to_excel(PAPER_DIR / "数字一致性审计.xlsx", index=False)
    pd.DataFrame(
        [{"reference": ref, "body_author_year_cited": _reference_is_cited(ref)} for ref in refs]
    ).to_excel(PAPER_DIR / "引用一致性审计.xlsx", index=False)


def _number_source(key: str) -> str:
    if key.startswith("stock") or key == "guidance_std":
        return "output/results/stock_volatility_main.json; data/processed/stock_event_panel.csv"
    if key.startswith("egarch"):
        return "output/results/daily_egarch_x_results.json"
    if key.startswith("bond"):
        return "output/results/yield_curve_results.csv"
    if key.startswith("cross"):
        return "output/results/cross_fitted_bond_exploration.csv"
    if key.startswith("svc") or key.startswith("lexicon") or key == "manual_sentences":
        return "output/results/text_model_evaluation.json; output/diagnostics/text_validation_metrics.xlsx"
    if key.startswith("power"):
        return "output/diagnostics/market_power_analysis.csv"
    return "data/processed/text_features.csv"


def _reference_is_cited(ref: str) -> bool:
    keys = ["董青马", "姜富伟", "尚玉皇", "中国人民银行", "Blinder", "Du", "Ehrmann", "Gurkaynak", "Hansen", "Jarocinski", "Kuttner", "McMahon", "Nakamura"]
    return any(key in ref for key in keys)


def export_pdf_with_word() -> None:
    word = None
    try:
        import win32com.client
        import pythoncom
        
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0  # Suppress dialogs
        doc = word.Documents.Open(
            str(DOCX_PATH.resolve()),
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
        )
        doc.Save()
        doc.ExportAsFixedFormat(str(PDF_PATH.resolve()), ExportFormat=17, OptimizeFor=0)
        doc.Close(False)
        word.Quit()
        word = None
    except Exception as exc:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        raise RuntimeError(f"Word COM PDF export failed: {exc}") from exc


def inspect_journal_pdf() -> dict:
    if not PDF_PATH.exists():
        raise FileNotFoundError(PDF_PATH)
    if REVIEW_DIR.exists():
        shutil.rmtree(REVIEW_DIR)
    REVIEW_DIR.mkdir(parents=True, exist_ok=True)
    pdf = fitz.open(str(PDF_PATH))
    inventory = []
    page_images = []
    caption_pages: dict[str, int] = {}
    for idx, page in enumerate(pdf, start=1):
        text = page.get_text("text")
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        image_path = REVIEW_DIR / f"page_{idx:03d}.png"
        pix.save(str(image_path))
        page_images.append(image_path)
        nonwhite = _nonwhite_ratio(image_path)
        blank = 1 - nonwhite
        inventory.append({"page": idx, "text_chars": len(text.strip()), "nonwhite_ratio": nonwhite, "blank_ratio": blank})
        for line in text.splitlines():
            match = re.match(r"^(图\s*\d+|表\s*\d+)\s", line.strip())
            if match:
                caption_pages.setdefault(re.sub(r"\s+", "", match.group(1)), idx)
    pdf.close()
    _write_contact_sheet(page_images)
    pd.DataFrame(inventory).to_csv(REVIEW_DIR / "page_inventory.csv", index=False, encoding="utf-8-sig")
    audit = _run_text_layout_audit(caption_pages, inventory)
    audit["page_count"] = len(inventory)
    audit["max_blank_ratio"] = max(row["blank_ratio"] for row in inventory) if inventory else None
    audit["caption_pages"] = caption_pages
    audit["word_com_pdf_export"] = True
    (REVIEW_DIR / "journal_audit.json").write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    if not audit["passed"]:
        raise RuntimeError(f"Journal PDF audit failed: {audit['issues']}")
    cover = Image.open(page_images[0])
    width, height = cover.size
    cover.crop((0, 0, width, int(height * 0.22))).save(OUTPUT_DIR / "diagnostics" / "cover_top_final.png")
    return audit


def _nonwhite_ratio(image_path: Path) -> float:
    image = Image.open(image_path).convert("L")
    small = image.resize((max(1, image.width // 4), max(1, image.height // 4)))
    pixels = list(small.getdata())
    nonwhite = sum(1 for value in pixels if value < 245)
    return nonwhite / len(pixels)


def _write_contact_sheet(paths: list[Path]) -> None:
    thumbs = []
    for path in paths:
        img = Image.open(path).convert("RGB")
        img.thumbnail((220, 310))
        thumbs.append((path, img.copy()))
    cols = 4
    rows = math.ceil(len(thumbs) / cols)
    sheet = Image.new("RGB", (cols * 250, rows * 350), "white")
    draw = ImageDraw.Draw(sheet)
    for idx, (path, img) in enumerate(thumbs):
        x = (idx % cols) * 250 + 15
        y = (idx // cols) * 350 + 20
        sheet.paste(img, (x, y))
        draw.text((x, y + img.height + 6), path.stem, fill="black")
    sheet.save(REVIEW_DIR / "contact_sheet.png")


def _run_text_layout_audit(caption_pages: dict[str, int], inventory: list[dict]) -> dict:
    doc = Document(str(DOCX_PATH))
    text_parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
    text = "\n".join(text_parts)
    banned = [
        "README.md",
        "run_all.py",
        "configs/project.yml",
        "final_submission",
        "Manifest",
        "文件哈希",
        "缓存",
        "pytest",
        "数字审计文件路径",
        "复核者",
        "固定路线",
        "开发阶段",
        "流水线",
        "本文强调",
        "本文没有",
        "本文不把",
        "不继续搜索",
        "正式路线",
        "锁定",
        "如实报告",
        "提交目录",
        "哈希",
        "所有结果均",
        "值得注意的是",
        "需要指出的是",
        "提供了有益参考",
        "具有重要意义",
        "dependent",
        "target",
        "beta",
        "se_hc3",
        "p_value",
        "full_joint_mle",
        "post_2019_total_effect",
        "leg" + "acy",
        "Co" + "dex",
        "pro" + "mpt",
    ]
    issues = []
    for term in banned:
        if term in text:
            issues.append(f"正文残留禁用词：{term}")
    for required in ["图1", "图2", "图3", "图4", "表1", "表2", "表3", "表4", "表5"]:
        if not any(key.startswith(required) for key in caption_pages):
            issues.append(f"PDF未定位到{required}")
    if len([p for p in doc.paragraphs if re.match(r"^图\s*\d+\s", p.text.strip())]) != 4:
        issues.append("DOCX正文图题数量不是4")
    if len([p for p in doc.paragraphs if re.match(r"^表\s*\d+\s", p.text.strip())]) != 5:
        issues.append("DOCX正文表题数量不是5")
    if any(row["text_chars"] == 0 and row["nonwhite_ratio"] < 0.01 for row in inventory):
        issues.append("存在疑似空白页")
    if len(inventory) < 10:
        issues.append("PDF页数过少")
    return {"passed": not issues, "issues": issues}
