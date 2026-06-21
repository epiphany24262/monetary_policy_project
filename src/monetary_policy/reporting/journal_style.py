from __future__ import annotations

from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PAGE_STYLE = {
    "width_cm": 21.0,
    "height_cm": 29.7,
    "top_cm": 2.3,
    "bottom_cm": 2.2,
    "left_cm": 2.5,
    "right_cm": 2.5,
}

FONT_STYLE = {
    "body_cn": "宋体",
    "body_en": "Times New Roman",
    "title_cn": "SimHei",
    "heading_cn": "仿宋",
    "subheading_cn": "黑体",
    "author_cn": "楷体",
    "abstract_cn": "仿宋",
    "reference_cn": "仿宋",
}


PARAGRAPH_STYLE = {
    "body_size_pt": 10.5,
    "body_line_pt": 18,
    "abstract_size_pt": 10.5,
    "abstract_line_pt": 15,
    "english_abstract_line_pt": 14,
    "first_line_chars": 2,
    "author_size_pt": 16,
}

HEADING_STYLE = {
    "level1_size_pt": 14,
    "level1_before_pt": 12,
    "level1_after_pt": 6,
    "level2_size_pt": 12,
    "level2_before_pt": 9,
    "level2_after_pt": 3,
}

TABLE_STYLE = {
    "caption_size_pt": 9,
    "body_size_pt": 9,
    "note_size_pt": 7.5,
    "top_border_sz": "8",
    "mid_border_sz": "4",
    "bottom_border_sz": "8",
}

FIGURE_STYLE = {
    "width_cm": 11.5,
    "height_cm": 4.5,
    "double_panel_width_cm": 12.0,
    "double_panel_height_cm": 6.2,
    "dpi": 300,
    "axis_font_pt": 8,
    "legend_font_pt": 7.5,
    "colors": ["#000000", "#4D4D4D", "#8C8C8C", "#BFBFBF"],
}

CAPTION_STYLE = {
    "size_pt": 9,
    "line_pt": 13,
}

HEADER_FOOTER_STYLE = {
    "header_text": "罗允绩：中国货币政策报告文本特征与金融市场反应",
    "header_size_pt": 9,
    "footer_size_pt": 9,
    "border_sz": "4",
}

REFERENCE_STYLE = {
    "title_size_pt": 12,
    "body_size_pt": 7.5,
    "line_pt": 13,
    "hanging_chars": 2,
}


def cm(value: float):
    return Cm(value)


def pt(value: float):
    return Pt(value)


def set_run_font(run, cn_font: str | None = None, size_pt: float | None = None, bold: bool = False, italic: bool = False) -> None:
    cn = cn_font or FONT_STYLE["body_cn"]
    run.font.name = FONT_STYLE["body_en"]
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), cn)
    r_fonts.set(qn("w:ascii"), FONT_STYLE["body_en"])
    r_fonts.set(qn("w:hAnsi"), FONT_STYLE["body_en"])
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def set_paragraph_format(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    line_pt: float | None = None,
    first_line_chars: float | None = None,
    before_pt: float = 0,
    after_pt: float = 0,
    exact: bool = True,
    keep_with_next: bool = False,
) -> None:
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    fmt.widow_control = True
    if keep_with_next:
        fmt.keep_with_next = True
        fmt.keep_together = True
    if line_pt is not None:
        if exact:
            fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            fmt.line_spacing = Pt(line_pt)
        else:
            fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            fmt.line_spacing = line_pt / 12.0
    if first_line_chars is not None:
        fmt.first_line_indent = Pt(PARAGRAPH_STYLE["body_size_pt"] * first_line_chars)


def set_figure_paragraph_format(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    before_pt: float = 4,
    after_pt: float = 1,
    keep_with_next: bool = True,
) -> None:
    """Format for paragraphs containing pictures — no exact line spacing."""
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    fmt.first_line_indent = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = 1.0
    if keep_with_next:
        fmt.keep_with_next = True
        fmt.keep_together = True


def configure_section(section) -> None:
    section.page_width = Cm(PAGE_STYLE["width_cm"])
    section.page_height = Cm(PAGE_STYLE["height_cm"])
    section.top_margin = Cm(PAGE_STYLE["top_cm"])
    section.bottom_margin = Cm(PAGE_STYLE["bottom_cm"])
    section.left_margin = Cm(PAGE_STYLE["left_cm"])
    section.right_margin = Cm(PAGE_STYLE["right_cm"])


def add_body_section(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_section(section)
    section.different_first_page_header_footer = True
    section.odd_and_even_pages_header_footer = True
    _restart_page_numbering(section, 1)
    section.first_page_header.is_linked_to_previous = False
    _clear_header_footer(section.first_page_header)
    section.first_page_footer.is_linked_to_previous = False
    _clear_header_footer(section.first_page_footer)
    section.footer.is_linked_to_previous = False
    _clear_header_footer(section.footer)
    try:
        if section.even_page_footer:
            section.even_page_footer.is_linked_to_previous = False
            _clear_header_footer(section.even_page_footer)
    except:
        pass
    _add_running_header(section)
    return section


def _clear_header_footer(part) -> None:
    for paragraph in part.paragraphs:
        paragraph.clear()


def _restart_page_numbering(section, start: int) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), str(start))
    sect_pr.append(pg_num)


def _add_running_header(section) -> None:
    def create_header_with_tabs(header, is_even):
        _clear_header_footer(header)
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.clear()
        
        pPr = para._element.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        # Center tab at ~8cm
        tab1 = OxmlElement('w:tab')
        tab1.set(qn('w:val'), 'center')
        tab1.set(qn('w:pos'), '4535') 
        tabs.append(tab1)
        # Right tab at ~16cm
        tab2 = OxmlElement('w:tab')
        tab2.set(qn('w:val'), 'right')
        tab2.set(qn('w:pos'), '9070') 
        tabs.append(tab2)
        pPr.append(tabs)
        
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4") # 0.5pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        p_bdr.append(bottom)
        pPr.append(p_bdr)
        
        def add_page_fld():
            run = para.add_run("· ")
            set_run_font(run, FONT_STYLE["body_en"], HEADER_FOOTER_STYLE["header_size_pt"])
            run_fld = para.add_run()
            set_run_font(run_fld, FONT_STYLE["body_en"], HEADER_FOOTER_STYLE["header_size_pt"])
            fld1 = OxmlElement('w:fldChar')
            fld1.set(qn('w:fldCharType'), 'begin')
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = "PAGE"
            fld2 = OxmlElement('w:fldChar')
            fld2.set(qn('w:fldCharType'), 'end')
            run_fld._r.append(fld1)
            run_fld._r.append(instr)
            run_fld._r.append(fld2)
            run2 = para.add_run(" ·")
            set_run_font(run2, FONT_STYLE["body_en"], HEADER_FOOTER_STYLE["header_size_pt"])
            
        def add_tab():
            run = para.add_run()
            tab = OxmlElement('w:tab')
            run._r.append(tab)

        if is_even:
            add_page_fld()
            add_tab()
            r_c = para.add_run("面向经济和金融的Python编程")
            set_run_font(r_c, FONT_STYLE["body_cn"], HEADER_FOOTER_STYLE["header_size_pt"])
            add_tab()
            r_r = para.add_run("2026年6月")
            set_run_font(r_r, FONT_STYLE["body_cn"], HEADER_FOOTER_STYLE["header_size_pt"])
        else:
            r_l = para.add_run("课程论文")
            set_run_font(r_l, FONT_STYLE["body_cn"], HEADER_FOOTER_STYLE["header_size_pt"])
            add_tab()
            r_c = para.add_run("罗允绩：中国货币政策报告文本特征与金融市场反应")
            set_run_font(r_c, FONT_STYLE["body_cn"], HEADER_FOOTER_STYLE["header_size_pt"])
            add_tab()
            add_page_fld()

    try:
        header = section.header
        header.is_linked_to_previous = False
        create_header_with_tabs(header, False)
    except Exception as e:
        pass

    try:
        even_header = section.even_page_header
        even_header.is_linked_to_previous = False
        create_header_with_tabs(even_header, True)
    except Exception as e:
        pass


def set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt: float | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(12)
    run = paragraph.add_run(str(text))
    set_run_font(run, FONT_STYLE["body_cn"], size_pt or TABLE_STYLE["body_size_pt"], bold=bold)


def set_table_width(table, widths_cm: list[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    if not widths_cm:
        return

    # Total expected: 15.0cm exactly = 8504 twips
    total_twips_target = 8504
    twips_per_cm = 1440 / 2.54
    
    col_twips = [int(w * twips_per_cm) for w in widths_cm]
    col_twips[-1] += total_twips_target - sum(col_twips)

    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(total_twips_target))

    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        table._tbl.remove(tblGrid)
    
    tblGrid = OxmlElement('w:tblGrid')
    for width_twips in col_twips:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width_twips))
        tblGrid.append(gridCol)
    # Ensure tblGrid is immediately after tblPr
    table._tbl.insert(table._tbl.index(tblPr) + 1, tblGrid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells[: len(col_twips)]):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(col_twips[idx]))

def set_cell_margins(cell, left_cm=0.10, right_cm=0.10, top_cm=0.0, bottom_cm=0.0) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    
    for edge, val_cm in [('top', top_cm), ('bottom', bottom_cm), ('left', left_cm), ('right', right_cm)]:
        el = tcMar.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tcMar.append(el)
        el.set(qn('w:type'), 'dxa')
        el.set(qn('w:w'), str(int(val_cm * 1440 / 2.54)))

def set_decimal_cell_text(cell, text: str, cell_width_twips: int, left_margin_twips: int = 57, right_margin_twips: int = 57, size_pt=None, bold=False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    
    # Text is primarily left aligned, with decimal tab pushing it to the right position
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(12)
    
    pPr = paragraph._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        pPr.append(tabs)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'decimal')
    
    # Calculate usable inner width and logical center
    usable_width = cell_width_twips - left_margin_twips - right_margin_twips
    # Because numbers display typically "-0.1124***", the right part is considerably wider.
    # We offset the anchor dynamically to ~45% from left margin to keep the block visually centered.
    anchor_pos = int(usable_width * 0.45)
    
    # For a decimal tab inside a cell, its position is relative to the cell's left margin.
    tab.set(qn('w:pos'), str(anchor_pos))
    tabs.append(tab)
    
    run_tab = paragraph.add_run('\t')
    run = paragraph.add_run(str(text))
    set_run_font(run, FONT_STYLE["body_en"], size_pt or TABLE_STYLE["body_size_pt"], bold=bold)


def mark_row_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def clear_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def set_cell_border(cell, edge: str, sz: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    element = tc_borders.find(qn(f"w:{edge}"))
    if element is None:
        element = OxmlElement(f"w:{edge}")
        tc_borders.append(element)
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), sz)
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), "000000")


def merge_row_cells(row):
    merged = row.cells[0]
    for cell in row.cells[1:]:
        merged = merged.merge(cell)
    return row.cells[0]
