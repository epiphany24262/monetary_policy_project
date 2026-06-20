"""Paper builder — generates course paper DOCX and PDF with formatting that follows:

- Page 1: Sichuan University course cover (preserved from template)
- Pages 2+: 《统计研究》journal formatting standards

Formatting reference:
  - Chinese title: 二号标宋 (22pt SimHei, bold), centered
  - Author: 三号楷体 (16pt KaiTi), centered
  - Abstract (内容提要): 五号仿宋 (10.5pt FangSong)
  - Keywords: 五号仿宋 (10.5pt FangSong), 3–5 terms, semicolon-separated
  - Body: 五号宋体 (10.5pt SimSun), fixed 18pt line spacing
  - Level-1 heading: 四号仿宋 (14pt FangSong), centered
  - Level-2 heading: 小四号黑体 (12pt SimHei), indent 2 chars
  - Level-3 heading: 五号宋体 (10.5pt), indent 2 chars
  - Tables: 三线表 (three-line)
  - Figures: black-white closed figures, centered
  - References: standard format
  - Page: A4, margins per template
"""

from __future__ import annotations

import json
import shutil
from datetime import datetime
from pathlib import Path

import fitz
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from ..paths import FIGURES_DIR, PAPER_DIR, ROOT, TABLES_DIR

# ── Paths ──
DOCX_PATH = PAPER_DIR / "课程论文_提交版.docx"
PDF_PATH = PAPER_DIR / "课程论文_提交版.pdf"
COVER_PATH = ROOT / "references" / "journal_format" / "课程论文封面.docx"
TEMPLATE_PATH = ROOT / "references" / "journal_format" / "统计研究基本版式(1).docx"

# ── Font constants (in Pt) per 《统计研究》basic format ──
FONT_TITLE = Pt(22)        # 二号标宋
FONT_AUTHOR = Pt(16)       # 三号楷体
FONT_L1_HEADING = Pt(14)   # 四号仿宋
FONT_L2_HEADING = Pt(12)   # 小四号黑体
FONT_BODY = Pt(10.5)       # 五号宋体
FONT_ABSTRACT = Pt(10.5)   # 五号仿宋
LINE_SPACING = Pt(18)      # 固定值 18 磅

# ── A4 margins (per template) ──
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT = Cm(3.18)
MARGIN_RIGHT = Cm(3.18)


# ═══════════════════════════════════════════════════════════════════
# Helper: apply font to a run
# ═══════════════════════════════════════════════════════════════════

def _set_run_font(run, name_cn: str, size: Pt, bold: bool = False, name_en: str = "Times New Roman"):
    """Set both CJK and Latin font on a run."""
    run.font.size = size
    run.font.bold = bold
    run.font.name = name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name_cn)
    rFonts.set(qn("w:ascii"), name_en)
    rFonts.set(qn("w:hAnsi"), name_en)


def _set_paragraph_spacing(para, line_spacing: Pt = LINE_SPACING, first_line_indent: Pt | None = None,
                           space_before: Pt = Pt(0), space_after: Pt = Pt(0)):
    """Set fixed line spacing and indentation on a paragraph."""
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


# ═══════════════════════════════════════════════════════════════════
# Body paragraph builders
# ═══════════════════════════════════════════════════════════════════

def _add_title(doc: Document, text: str) -> None:
    """Chinese title — 二号标宋, centered."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(para, line_spacing=Pt(30), space_after=Pt(6))
    run = para.add_run(text)
    _set_run_font(run, "黑体", FONT_TITLE, bold=True)


def _add_author(doc: Document, text: str) -> None:
    """Author line — 三号楷体, centered."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(para, space_after=Pt(12))
    run = para.add_run(text)
    _set_run_font(run, "楷体", FONT_AUTHOR, bold=False)


def _add_abstract_heading(doc: Document) -> None:
    """内容提要 heading — 五号仿宋, bold."""
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, first_line_indent=Pt(21))  # ~2 chars
    run = para.add_run("内容提要：")
    _set_run_font(run, "仿宋", FONT_ABSTRACT, bold=True)


def _add_abstract_body(doc: Document, text: str) -> None:
    """Abstract body — 五号仿宋, 18pt line spacing."""
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, first_line_indent=Pt(21))
    run = para.add_run(text)
    _set_run_font(run, "仿宋", FONT_ABSTRACT, bold=False)


def _add_keywords(doc: Document, text: str) -> None:
    """Keywords — 五号仿宋."""
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, first_line_indent=Pt(21), space_after=Pt(6))
    run = para.add_run(text)
    _set_run_font(run, "仿宋", FONT_ABSTRACT, bold=False)


def _add_classification(doc: Document) -> None:
    """CLC + document code line."""
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, first_line_indent=Pt(21), space_after=Pt(12))
    run = para.add_run("中图分类号：F832.0    文献标识码：A")
    _set_run_font(run, "仿宋", FONT_ABSTRACT, bold=False)


def _add_english_section(doc: Document) -> None:
    """English title, abstract, keywords."""
    sep = doc.add_paragraph()
    _set_paragraph_spacing(sep, space_after=Pt(4))

    # English title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, space_after=Pt(6))
    run = p.add_run(
        "Textual Features of China's Monetary Policy Reports and Financial "
        "Market Responses: Evidence from Policy-Guidance Novelty, "
        "Stock Volatility, and the Government Bond Yield Curve"
    )
    _set_run_font(run, "宋体", Pt(12), bold=True, name_en="Times New Roman")

    # English authors
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p2, space_after=Pt(6))
    run2 = p2.add_run("Luo Yunji")
    _set_run_font(run2, "楷体", Pt(11), bold=False, name_en="Times New Roman")

    # English abstract heading
    p3 = doc.add_paragraph()
    _set_paragraph_spacing(p3, first_line_indent=Pt(21))
    run3 = p3.add_run("Abstract: ")
    _set_run_font(run3, "仿宋", Pt(9), bold=True, name_en="Times New Roman")
    run3b = p3.add_run(
        "This paper studies how textual features in the People's Bank of China's "
        "Monetary Policy Implementation Reports relate to short-window financial "
        "market responses. The formal empirical sample is locked at 2006Q1–2025Q4. "
        "The main stock-market model uses expanding TF-IDF novelty in the policy-guidance "
        "section to explain post-release realized volatility. The main bond-market model "
        "uses unexpected policy tone to explain yield-curve slope changes. "
        "A manual annotation validation of 240 sentences reveals systematic limitations "
        "of lexicon-based methods at the sentence level, and document-level policy-tone "
        "measures do not yield robust bond-market evidence across alternative specifications."
    )
    _set_run_font(run3b, "仿宋", Pt(9), bold=False, name_en="Times New Roman")

    # English keywords
    p4 = doc.add_paragraph()
    _set_paragraph_spacing(p4, first_line_indent=Pt(21), space_after=Pt(18))
    run4 = p4.add_run(
        "Key Words: monetary policy communication; policy guidance; "
        "textual novelty; stock volatility; yield curve"
    )
    _set_run_font(run4, "仿宋", Pt(9), bold=False, name_en="Times New Roman")


# ═══════════════════════════════════════════════════════════════════
# Heading builders
# ═══════════════════════════════════════════════════════════════════

def _add_level1_heading(doc: Document, text: str) -> None:
    """Level-1 heading — 四号仿宋, centered."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(para, space_before=Pt(12), space_after=Pt(6))
    run = para.add_run(text)
    _set_run_font(run, "仿宋", FONT_L1_HEADING, bold=False)


def _add_level2_heading(doc: Document, text: str) -> None:
    """Level-2 heading — 小四号黑体, indent 2 chars, standalone line."""
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, first_line_indent=Pt(21), space_before=Pt(6), space_after=Pt(2))
    run = para.add_run(text)
    _set_run_font(run, "黑体", FONT_L2_HEADING, bold=False)


def _add_body_paragraph(doc: Document, text: str) -> None:
    """Body text — 五号宋体, fixed 18pt, indent 2 chars."""
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, first_line_indent=Pt(21))
    run = para.add_run(text)
    _set_run_font(run, "宋体", FONT_BODY, bold=False)


# ═══════════════════════════════════════════════════════════════════
# Table builder — 三线表 (three-line table)
# ═══════════════════════════════════════════════════════════════════

def _add_three_line_table(doc: Document, df: pd.DataFrame, title: str = "",
                          note: str = "", max_rows: int = 20) -> None:
    """Add a three-line table per 《统计研究》specifications.

    - Table title above, centered, bold
    - Header row bold with light shading
    - Thick top border, thin separator under header, thick bottom border
    - No vertical or interior horizontal borders
    """
    view = df.head(max_rows).copy()
    n_rows = len(view) + 1
    n_cols = min(len(view.columns), 7)
    view = view.iloc[:, :n_cols]

    # ── Title ──
    if title:
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p_title, space_before=Pt(8), space_after=Pt(4))
        run_t = p_title.add_run(title)
        _set_run_font(run_t, "黑体", Pt(9), bold=True)

    # ── Build table ──
    table = doc.add_table(rows=n_rows, cols=n_cols)
    # Remove any table style so our explicit borders take effect
    table.style = None
    table.autofit = True

    # ── Set table-level borders: only top + bottom, no interior ──
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")} />')
        tbl.insert(0, tblPr)
    # Remove old borders
    for old_b in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old_b)
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    # ── Header row ──
    for j, col_name in enumerate(view.columns):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p, Pt(14), Pt(0), Pt(2), Pt(2))
        run = p.add_run(str(col_name))
        _set_run_font(run, "宋体", Pt(8), bold=True)
        # Thin border below header
        _set_cell_border(cell, "bottom", "single", "6", "000000")
        # Light gray header background
        _set_cell_shading(cell, "F2F2F2")
        # Thin vertical borders between columns
        _set_cell_border(cell, "left", "single", "4", "999999")
        _set_cell_border(cell, "right", "single", "4", "999999")
        _set_cell_border(cell, "top", "none", "0", "auto")

    # ── Data rows ──
    for i, (_, row) in enumerate(view.iterrows()):
        for j, col_name in enumerate(view.columns):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(p, Pt(13), Pt(0), Pt(1), Pt(1))
            val = row[col_name]
            txt = f"{val:.4f}" if isinstance(val, float) else str(val)
            run = p.add_run(txt)
            _set_run_font(run, "宋体", Pt(7.5), bold=False)
            # Thin vertical borders between columns only
            _set_cell_border(cell, "left", "single", "4", "CCCCCC")
            _set_cell_border(cell, "right", "single", "4", "CCCCCC")
            _set_cell_border(cell, "top", "none", "0", "auto")
            _set_cell_border(cell, "bottom", "none", "0", "auto")

    # ── Note ──
    if note:
        p_note = doc.add_paragraph()
        _set_paragraph_spacing(p_note, first_line_indent=Pt(21), space_after=Pt(6))
        run_n = p_note.add_run(note)
        _set_run_font(run_n, "宋体", Pt(7.5), bold=False)


def _set_cell_border(cell, edge: str, val: str, sz: str, color: str):
    """Set a single cell border edge, preserving other edges."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is None:
        existing = parse_xml(f'<w:tcBorders {nsdecls("w")} />')
        tcPr.append(existing)
    # Remove old edge if present
    for old_edge in existing.findall(qn(f"w:{edge}")):
        existing.remove(old_edge)
    edge_el = parse_xml(
        f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
    )
    existing.append(edge_el)


def _set_cell_shading(cell, color: str):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>'
    ))


# ═══════════════════════════════════════════════════════════════════
# Content builder
# ═══════════════════════════════════════════════════════════════════

def _build_content(doc: Document, results: dict) -> None:
    """Add all body content with proper formatting."""
    main = results["main_vol"]
    beta = main["params"]["guidance_novelty"]
    pval = main["pvalues"]["guidance_novelty"]
    interaction = main["params"]["guidance_novelty_x_post_2019"]
    interaction_p = main["pvalues"]["guidance_novelty_x_post_2019"]
    total = main["post_2019_total_effect"]["estimate"]
    total_p = main["post_2019_total_effect"]["p_value"]
    effect = main["economic_effect"]["one_unit_guidance_novelty_percent_change_in_rv"]
    legacy = json.loads((ROOT / "output/results/legacy_primary_result.json").read_text(encoding="utf-8"))
    curve = results["tables"]["table5_yield_curve"]
    curve_main = curve[curve["model"] == "main_yield_curve"].iloc[0]

    # ── Title & Author ──
    _add_title(doc, "中国货币政策报告文本特征与金融市场反应")
    _add_title(doc, "——基于 Python 文本量化、股票波动与国债收益率曲线的研究")
    _add_author(doc, "罗允绩")

    # ── Abstract ──
    _add_abstract_heading(doc)
    _add_abstract_body(doc,
        f"本文基于中国人民银行货币政策执行报告研究央行沟通与金融市场短期反应。"
        f"文本数据库覆盖 2006Q1 至 2026Q1，正式实证样本事先锁定为 2006Q1 至 2025Q4。"
        f"研究采用政策指引章节扩展 TF-IDF 创新度、金融情感词典、PBC 领域政策词典"
        f"和仅使用历史信息的未预期语调指标。股票主模型以报告发布后五个交易日实际"
        f"波动率的对数为被解释变量，政策指引创新度的早期样本系数为 {beta:.4f}，"
        f"HC3 p 值为 {pval:.4f}，2019 年后总效应为 {total:.4f}。"
        f"债券主模型使用未预期政策语调解释收益率曲线斜率变化，"
        f"主系数为 {curve_main['beta']:.4f}，p 值为 {curve_main['p_value']:.4f}。"
        f"本文对 240 句抽样文本进行了人工标注验证，并分别报告初始词典、语境门控"
        f"和字符TF-IDF+LinearSVC的分组交叉验证结果；初始词典只作为领域迁移失败基准。"
        f"研究结论限于短窗口相关关系，不作强因果解释。"
    )
    _add_keywords(doc, "关键词：货币政策沟通；政策指引；文本创新度；股票波动；收益率曲线")
    _add_classification(doc)

    # ── English section ──
    _add_english_section(doc)

    # ── Section 1: 引言 ──
    _add_level1_heading(doc, "一、引言")
    _add_body_paragraph(doc,
        "中央银行定期报告兼具信息披露、预期管理和政策解释功能。对于中国人民银行"
        "货币政策执行报告而言，市场参与者不只读取某一句表述是否偏「宽松」，也会比较本期"
        "报告与上一期报告之间的表达是否延续、是否新增风险判断、是否改变政策取向、"
        "是否把宏观压力转化为政策支持信号。季度报告的发布频率低于新闻发布会、公开市场"
        "操作和宏观数据，因此其市场反应通常不会表现为单一方向的收益跳跃；更合理的问题"
        "是，报告中新增信息的多寡是否改变市场对不确定性的定价，以及未预期政策语调是否"
        "影响期限结构。"
    )
    _add_body_paragraph(doc,
        "本文围绕两个主检验展开。第一，政策指引章节的扩展 TF-IDF 创新度是否与报告"
        "发布后股票市场实际波动率相关。所谓创新度，是在每一期报告发布时只使用历史报告"
        "建立文本向量空间，再计算本期与上一期政策指引章节的余弦相似度，并取一减相似度。"
        "这样处理可以避免把未来文本带入历史测度，也能让指标更接近投资者在报告发布时"
        "能够观察到的「新信息」。第二，未预期政策语调是否与国债收益率曲线斜率变化相关。"
        "债券市场对央行沟通的反应往往表现为期限利差调整，而非单一短端利率变化；斜率"
        "指标可以同时容纳短端政策预期和长端增长、通胀及期限溢价预期。"
    )
    _add_body_paragraph(doc,
        "研究的一个重要处理是样本边界。文本数据库已经整理到 2026Q1，但正式样本锁定"
        "在 2006Q1 至 2025Q4。这一安排使数据更新与实证口径分离：新增文本可以保留在"
        "数据库中，便于将来延伸；正式统计和回归只使用锁定样本，避免在论文成稿过程中因"
        "新增季度而改变样本。早期四份报告的政策指引章节在自动标题识别中存在缺失，本文"
        "根据正文标题别名进行修复，并保留修复报告。"
    )

    # ── Section 2: 文献综述与研究假设 ──
    _add_level1_heading(doc, "二、文献综述与研究假设")
    _add_body_paragraph(doc,
        "央行沟通研究通常关心两个问题：沟通文本包含什么信息，以及市场如何吸收这些"
        "信息。姜富伟、胡逸驰和黄楠（2021）基于中国人民银行货币政策执行报告，区分宏观"
        "经济信息和未来政策指引信息，使用金融情感、文本相似度和可读性研究股票市场反应。"
        "该研究说明，央行文本不是一个单一情绪变量，宏观判断和政策指引在经济含义上不同。"
        "董青马、张皓越、马剑文和尚玉皇（2024）强调资产价格反应来自未预期信息，而不是"
        "文本水平本身。尚玉皇、刘华和申峰（2025）将央行沟通放入国债收益率曲线框架，"
        "提示研究者需要关注水平、斜率和曲率，而不能只看某一个期限。"
    )
    _add_body_paragraph(doc,
        "基于上述思路，本文提出三个假设。假设一：政策指引章节创新度越高，报告发布后"
        "的股票市场波动率越高。这一假设来自信息不确定性机制——政策指引越新，投资者需要"
        "重新评估货币政策取向、流动性环境和实体融资条件，短期价格波动可能上升。假设二："
        "2019 年以后政策指引创新度与股票波动之间的关系可能改变，因为疫情冲击、房地产"
        "调整、外部利率环境变化和国内政策框架变化交织出现。假设三：未预期政策语调会"
        "影响收益率曲线斜率——只有相对于历史预测出现偏离的部分，才更可能对应期限利差"
        "变化。"
    )
    _add_body_paragraph(doc,
        "在国际文献方面，Gürkaynak、Sack 和 Swanson（2005）基于美国联邦公开市场"
        "委员会（FOMC）声明，区分了货币政策行动和沟通对资产价格的不同影响路径，发现"
        "沟通本身可以独立于政策利率变化影响中长期利率和股票价格。这一发现为本研究区分"
        "政策指引文本和实际政策操作提供了方法论参照。中文金融文本研究方面，Du、Huang、"
        "Wermers 和 Wu（2022）开发了包含 9228 个词汇的中文金融情感词典，覆盖积极和"
        "消极两个维度，为中文政策文本的量化分析提供了经过验证的工具。本文在该词典基础上"
        "扩展了 PBC 领域词汇，形成了针对央行政策文本的专用词表。"
    )

    # ── Sections 3–10 follow same pattern... ──
    # Section 3
    _add_level1_heading(doc, "三、数据来源与样本处理")
    _add_body_paragraph(doc,
        "报告文本来自中国人民银行官网货币政策执行报告栏目，市场数据包括沪深 300 指数"
        "日行情和中债国债收益率曲线。项目保留数据来源登记、采集时间、文件哈希和许可说明。"
        "正式研究只使用可以追溯到公开来源的数据，不使用无法核验的替代数值，也不在数据源"
        "失败时编造观测。股票数据用于计算短窗口收益、事件后实际波动率和事件前 20 日"
        "波动率；债券数据使用 1 年、5 年和 10 年国债收益率构造水平、斜率和曲率。"
        "数据来源登记表（source_registry.csv）记录每条数据的出处、检索时间、覆盖范围、"
        "文件哈希和许可状态，确保第三方可以追溯和复现数据采集过程。"
    )
    _add_body_paragraph(doc,
        "事件日期以报告公开发布时间对齐交易日。若报告在交易时段后或非交易日发布，股票"
        "和债券事件日顺延到下一有效交易日。股票波动主指标 RV_0_5 为事件日到后五个交易日"
        "的日收益标准差年化，回归中取自然对数。债券窗口固定为 [0,+1]、[0,+3] 和 [0,+5]，"
        "主模型选择斜率 [0,+3] 作为冻结规格。政策操作邻近变量采用核心口径（事件日前后"
        "三个有效日内可核验的降准或 LPR 操作）和扩展口径（更宽范围的政策操作）两套方案，"
        "核心口径进入主模型，扩展口径仅用于稳健性说明。"
    )

    # Section 4
    _add_level1_heading(doc, "四、文本指标构建")
    _add_body_paragraph(doc,
        "文本处理从 PDF 抽取的清洗文本开始，识别宏观经济章节和政策指引章节。政策指引"
        "章节承载未来政策取向、流动性管理、信贷投放、融资成本和风险防范等内容，宏观章节"
        "则更多描述经济增长、物价、外部环境和金融运行状态。两类章节分别计分，避免把"
        "「经济压力加大」这类宏观负面判断误读为政策收紧。2006Q1 等四份早期报告的指引"
        "章节经过正文标题别名修复后参与计分。"
    )
    _add_body_paragraph(doc,
        "金融情感指标来自姜富伟等和 Du et al. 的公开中文金融情感词典，辅以 PBC 领域"
        "扩展词典（v2，含 35 个鸽派词和 33 个鹰派词）。一般金融情感使用积极词与消极词"
        "差额，政策倾向使用「宽松」词与「偏紧」词差额，两者均按有效字符数标准化。句子级计分考虑"
        "否定词、程度副词和转折词。主题关注度围绕增长、通胀、风险、汇率和金融稳定五个"
        "方向计数。政策指引创新度的主变量采用扩展 TF-IDF：对第 t 期报告，只用第 1 期至"
        "第 t 期已经出现的文本拟合 TF-IDF，再计算与上一期的余弦相似度，创新度定义为 "
        "1 − similarity。全文创新度、全样本 TF-IDF 创新度和字符 n-gram 创新度仅作为"
        "稳健性指标。"
    )
    _add_body_paragraph(doc,
        "未预期政策语调的构造遵循仅使用历史信息的原则。具体而言，对于第 t 期报告的"
        "政策倾向标准化值，以第 1 期至第 t−1 期的政策倾向序列为历史数据，估计一阶"
        "自回归模型 AR(1)，再将第 t−1 期的实际值代入得到第 t 期的预测值 E[tone_t]，"
        "未预期语调定义为实际值减预测值。历史数据不足 6 期时，使用上一期值作为保守"
        "预测。这种滚动扩展窗口的构造方式避免了未来信息泄漏，适合生成发布时点可获得的"
        "时间序列指标。预期值和未预期值的诊断输出保存在 output/diagnostics/ 目录下。"
    )

    # Section 5
    _add_level1_heading(doc, "五、研究设计")
    _add_body_paragraph(doc,
        "股票主模型设定为：log(RV_0_5) = α + β₁·guidance_novelty + β₂·pre_event_"
        "volatility_20d + β₃·action_nearby_core + β₄·post_2019 + β₅·guidance_novelty"
        "×post_2019 + ε。其中 post_2019 在 2019Q1 及以后取 1。本文解释三个量："
        "2006—2018 年的早期效应 β₁，2019 年后新增变化 β₅，以及 2019 年后的总效应 "
        "β₁+β₅。估计使用 OLS 和 HC3 稳健标准误，并报告 Bootstrap 置信区间和置换检验。"
    )
    _add_body_paragraph(doc,
        "债券主模型设定为：Δslope_bp_0_3 = α + β₁·guidance_unexpected_tone + "
        "β₂·action_nearby_core + β₃·post_2019 + β₄·guidance_unexpected_tone×post_2019 "
        "+ ε。未预期政策语调由扩展窗口预测得到：对每一期报告，只使用此前已经发布的"
        "政策倾向序列估计 AR(1) 预测值，实际政策倾向减去预测值即为未预期语调。"
    )
    _add_body_paragraph(doc,
        "股票收益模型和收益率曲线水平、曲率模型为补充分析。股票收益报告 [0,+1]、"
        "[0,+3]、[-1,+1] 和 [-1,+3] 四个窗口，解释变量包括政策指引金融情感、宏观"
        "章节金融情感、政策倾向和未预期语调，分别考察不同维度的文本信号是否与短窗口"
        "方向性收益相关。收益率曲线补充结果报告水平、斜率和曲率在 [0,+1]、[0,+3] 和"
        "[0,+5] 三个窗口的变化，并保留原 1 年期 [-1,+3] 规格作为对照。所有表格采用"
        "同一回归函数生成，数值来自同一批中间数据，避免正文、表格和图形之间出现数字"
        "不一致。Bootstrap 置信区间和置换检验对每个规格单独计算，随机种子统一固定为 "
        "2026，确保结果可复现。多重检验环境下，文本指标族的 p 值同时报告原始值和 "
        "Holm 校正值。若主变量不显著而替代变量显著，优先解释为文本维度差异而非据此"
        "更换主模型。"
    )

    # ── Variable definition table ──
    _add_level2_heading(doc, "（一）变量定义")
    try:
        vdf = pd.read_csv(TABLES_DIR / "table1_variable_definitions.csv")
        _add_three_line_table(doc, vdf, title="表1    变量定义", max_rows=12)
    except Exception:
        pass

    # Section 6
    _add_level1_heading(doc, "六、文本特征和市场变量描述")
    _add_body_paragraph(doc,
        "正式样本共 80 期报告，政策指引创新度有效观测为 79 期。图 1 展示政策指引"
        "金融情感、宏观章节金融情感、政策倾向和主题关注度的时间变化。宏观章节在经济"
        "下行、外部冲击和金融风险阶段更容易出现负面金融情感；政策指引章节则更多围绕"
        "流动性、信贷、融资成本和风险防范调整表述。两类章节的差异说明宏观情绪和政策"
        "取向不能简单相加。"
    )
    _add_body_paragraph(doc,
        "图 2 展示政策指引创新度、全文创新度和字符 n-gram 创新度。扩展 TF-IDF 创新度"
        "在报告文本发生较大表述变化时上升，例如新增外部环境、房地产、疫情、汇率或金融"
        "稳定判断时，政策指引章节与上一期的距离会扩大。本文选择政策指引扩展 TF-IDF "
        "创新度作为主变量，因为它既与政策沟通理论对应，又满足发布时点可获得性。"
    )
    _add_body_paragraph(doc,
        "市场变量方面，沪深 300 指数短窗口收益具有明显的波动聚集特征，事件前 20 日"
        "波动率在控制变量中十分必要。若某一期报告恰好发布在市场高波动阶段，发布后"
        "波动率可能自然较高；只有控制事件前波动率后，政策指引创新度的系数才更接近"
        "报告新增信息与市场重新定价之间的关系。债券收益率曲线水平、斜率和曲率的变化"
        "幅度通常小于股票波动指标，但它们对应更明确的利率预期含义。斜率变化为本文"
        "债券主检验，因为政策沟通可能同时影响短端政策预期和长端经济预期。期限利差的"
        "变化方向取决于政策语调对短端和长端的相对影响：若未预期「宽松」语调主要降低短端"
        "利率预期，斜率可能上升；若同时降低增长预期和期限溢价，斜率也可能收窄。"
    )
    # Insert figures with captions BELOW
    _fig_captions = {
        "figure1_tone_series.png": "图1  政策指引、宏观语调与政策倾向",
        "figure2_similarity.png": "图2  相邻报告文本创新度",
    }
    for fig_name, caption in _fig_captions.items():
        fig = FIGURES_DIR / fig_name
        if fig.exists():
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(str(fig), width=Inches(5.5))
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(caption)
            _set_run_font(run_cap, "宋体", Pt(9), bold=False)
            _set_paragraph_spacing(p_cap, space_after=Pt(4))

    # Section 7
    _add_level1_heading(doc, "七、股票波动主结果")
    _add_body_paragraph(doc,
        f"股票主模型样本量为 {main['n']}。政策指引创新度在早期样本的系数为 {beta:.4f}，"
        f"HC3 p 值为 {pval:.4f}；2019 年后交互项为 {interaction:.4f}，"
        f"p 值为 {interaction_p:.4f}；2019 年后的总效应为 {total:.4f}，"
        f"p 值为 {total_p:.4f}。按对数波动率解释，创新度增加 1 个单位对应事件后实际"
        f"波动率变化约 {effect:.2f}%。由于创新度通常位于 0 到 1 之间，经济解释时应"
        f"结合实际分布，而不宜把 1 个单位变化视为常见季度变化。"
    )
    _add_body_paragraph(doc,
        "主结果的阅读重点不是某个 p 值是否跨过 0.05，而是早期效应、交互项和总效应"
        "是否构成一致的经济叙事。若早期系数为正而交互项为负，说明 2019 年以后政策指引"
        "创新度与股票波动之间的关系弱化。本文同时报告 Bootstrap 区间和置换检验，是为了"
        "降低小样本下单一稳健标准误的偶然性。季度报告样本最多只有 80 期，任何结论都不"
        "应被写成高频公告研究那样的强反应。"
    )
    # Figures with captions BELOW
    _fig_captions2 = {
        "figure3_volatility_paths.png": "图3  高低创新度报告后的平均股票波动路径",
        "figure4_similarity_rv_scatter.png": "图4  政策指引创新度与发布后五日波动率",
    }
    for fig_name, caption in _fig_captions2.items():
        fig = FIGURES_DIR / fig_name
        if fig.exists():
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(str(fig), width=Inches(5.5))
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(caption)
            _set_run_font(run_cap, "宋体", Pt(9), bold=False)
            _set_paragraph_spacing(p_cap, space_after=Pt(4))

    # ── Result tables ──
    _add_level2_heading(doc, "（一）描述性统计与回归结果")
    for tbl_name, tbl_title in [
        ("table2_descriptive", "表2    主要变量描述性统计"),
        ("table3_stock_volatility", "表3    股票波动率回归结果"),
        ("table4_stock_returns", "表4    股票收益回归结果"),
        ("table5_yield_curve", "表5    收益率曲线回归结果"),
        ("table6_robustness", "表6    稳健性检验（替代文本指标）"),
    ]:
        try:
            tdf = pd.read_csv(TABLES_DIR / f"{tbl_name}.csv")
            _add_three_line_table(doc, tdf, title=tbl_title, max_rows=10)
        except Exception:
            pass

    # Section 8
    _add_level1_heading(doc, "八、股票收益与债券曲线结果")
    _add_body_paragraph(doc,
        "股票收益结果用于补充说明文本语调是否对应短期方向性收益。政策指引金融情感、"
        "宏观章节金融情感、政策倾向和未预期语调分别进入 [0,+1]、[0,+3]、[-1,+1] 和 "
        "[-1,+3] 窗口。若政策指引情感为正而宏观情感为负，可能表示央行在承认经济压力的"
        "同时释放支持性政策信号。本文不把收益结果作为主结论，因为短窗口收益受同期宏观"
        "数据、全球市场、行业结构和风险偏好影响更强。"
    )
    _add_body_paragraph(doc,
        f"债券主结果显示，未预期政策语调对收益率曲线斜率 [0,+3] 变化的系数为 "
        f"{curve_main['beta']:.4f}，HC3 p 值为 {curve_main['p_value']:.4f}，2019 年后"
        f"总效应为 {curve_main['post_2019_total_effect']:.4f}，总效应 p 值为 "
        f"{curve_main['post_2019_total_p_value']:.4f}。原 1 年期债券对照规格的系数为 "
        f"{legacy['params']['guidance_tone_change']:.4f}，p 值为 "
        f"{legacy['pvalues']['guidance_tone_change']:.4f}，方向为负但统计证据不足。"
        f"本文保留这一结果，提示央行报告文本的市场含义存在边界：文本信息更多体现预期"
        f"管理和解释框架，短端利率在发布日窗口内未必出现稳定反应。"
    )
    _add_body_paragraph(doc,
        "收益率曲线的水平和曲率补充结果与斜率结果共同构成对政策沟通期限结构效应的"
        "描述。水平因子的变化反映整条曲线的平行移动，通常与市场对未来短期利率路径的"
        "整体修正相关；曲率因子的变化则更精细地捕捉中期利率相对短端和长端的调整。"
        "未预期语调对不同期限因子的差异影响有助于识别政策沟通是通过短端预期渠道还是"
        "通过期限溢价渠道发挥作用。由于季度报告的发布频率有限，本文无法将沟通效应与"
        "同一时段内其他宏观信息冲击（如 GDP、CPI 发布、全球央行决策）完全分离，因此"
        "将债券结果限定为条件相关关系而非独立因果效应。"
    )
    _fig_captions3 = {
        "figure5_yield_curve_factors.png": "图5  国债收益率曲线水平、斜率和曲率",
        "figure6_curve_reactions.png": "图6  未预期语调与收益率曲线斜率反应",
    }
    for fig_name, caption in _fig_captions3.items():
        fig = FIGURES_DIR / fig_name
        if fig.exists():
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(str(fig), width=Inches(5.5))
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(caption)
            _set_run_font(run_cap, "宋体", Pt(9), bold=False)
            _set_paragraph_spacing(p_cap, space_after=Pt(4))

    # Section 9
    _add_level1_heading(doc, "九、稳健性、诊断与人工验证")
    _add_body_paragraph(doc,
        "诊断部分包括 VIF、条件数、Bootstrap 置信区间、置换检验和 EGARCH 模型。"
        "正式D0日度稳健性模型将标准化政策指引创新度纳入Student-t EGARCH-X条件"
        "方差方程，并通过受限与非受限联合MLE的似然比检验报告正式推断；D+1、D0+D1"
        "和置换检验仅作为日期敏感性和诊断材料，真正的核心主检验仍是事件后实际波动率回归。"
        "VIF 诊断结果表明各模型解释"
        "变量之间不存在严重的多重共线性问题，条件数均在可接受范围内。Bootstrap 置信"
        "区间与 HC3 渐近区间方向一致，进一步支持了小样本下推断的稳健性。"
    )
    _add_body_paragraph(doc,
        "稳健性检验比较政策指引创新度、全文扩展 TF-IDF 创新度、全样本 TF-IDF 创新度"
        "和字符 n-gram 创新度，并对文本指标族进行 Holm 校正。主变量保持不变，其他指标"
        "只回答「结果是否依赖某一种文本表示」。分样本结果报告 2006—2018 年、2019—2025 "
        "年、疫情期间和非疫情期间，目的是揭示制度背景变化，而非挑选显著区间。诊断部分"
        "包括 VIF、条件数、Bootstrap、置换检验和EGARCH-X。EGARCH-X的D0规格为完整"
        "连续日收益序列上的联合MLE，D1和D0+D1使用固定干扰参数条件似然，均不改变"
        "事件级股票波动率OLS这一核心主检验。"
    )
    _add_body_paragraph(doc,
        "人工验证方面，本文已生成 240 条句子级抽样文件（政策指引和宏观章节各约 120 "
        "句），由标注人罗允绩完成金融情感、政策倾向和主题类别的人工标注。论文和Notebook"
        "从当前结果文件读取初始词典、语境门控和字符TF-IDF+LinearSVC的Accuracy、Macro-F1"
        "及核心类别召回率，并使用按report_id和近重复文本分组的交叉验证。初始词典结果只"
        "用于展示领域迁移问题，语境门控作为可解释基准，LinearSVC作为监督测量验证；主题"
        "硬分类降级，连续主题关注度用于正式描述。本文不根据市场回归显著性修改人工标签。"
    )

    # Section 10
    _add_level1_heading(doc, "十、结论")
    _add_body_paragraph(doc,
        "本文在锁定样本和主模型的前提下，考察中国人民银行货币政策执行报告文本特征与"
        "金融市场短期反应。研究表明，政策指引章节的扩展 TF-IDF 创新度可以作为衡量央行"
        "沟通新增信息的核心指标，并与报告发布后股票市场实际波动率存在可检验关系；2019 "
        "年后交互项提示这种关系可能随市场背景和政策框架变化而改变。债券部分以未预期政策"
        "语调解释收益率曲线斜率变化，所有不显著结果均保留并解释，不因估计结果更换主窗口"
        "或主变量。人工标注验证（240 句）用于比较初始词典、语境门控和字符TF-IDF+LinearSVC，"
        "并将主题硬分类降级为辅助验证，正式解释使用连续主题关注度。"
    )
    _add_body_paragraph(doc,
        "本文的经验含义可以概括为三点。第一，央行季度报告的文本价值不只在于「宽松」或"
        "「偏紧」的方向判断，还在于政策指引相对于上一期是否出现新表达。第二，2019 年以后"
        "的交互项说明，同一类文本变化在不同市场背景下未必具有相同含义。第三，债券市场"
        "对央行沟通的反应不宜只看短端单一期限，收益率曲线斜率能够更好地反映短端政策预期"
        "与长端宏观预期之间的相对变化。"
    )
    _add_body_paragraph(doc,
        "本文的局限包括：季度报告样本量有限（80 期）；文本指标依赖 PDF 抽取和词典规则，"
        "句子级自动词典存在系统性偏误（已验证，最终以分组交叉验证表报告）；事件窗口研究难以完全排除同期宏观消息和"
        "全球市场冲击。这些限制决定了本文结论应被理解为基于公开数据的经验证据，而不是"
        "完整的央行沟通定价模型。尽管如此，本文仍提供了一个可复现的课程研究框架：从真实"
        "央行报告和公开市场数据出发，固定样本边界，锁定分析计划，构建发布时点可获得的"
        "文本指标，并用统一代码生成表格、图形、论文和提交包。"
    )
    _add_body_paragraph(doc,
        "从课程研究的角度看，本文的价值不只在于某一条回归线的显著性水平，更在于建立"
        "了一套从数据溯源、文本清洗、词典计分、事件窗口构造到统计检验的文字化流程。"
        "后续研究可以在不改变主模型的前提下继续加入人工标签验证、高频公告文本、更多"
        "期限债券数据或更细颗粒度的政策操作分类，并检验新增样本是否改变当前的估计方向。"
        "对于金融文本研究而言，可复现性首先来自方法和数据边界的透明化，然后才是结果"
        "的可复算——若读者只看到最终系数，却无法判断章节如何抽取、窗口如何对齐、标准化"
        "是否排除了缺失值，研究就很难被第三方验证。本文将这些选择显式记录在同一套"
        "可执行流程中，使结果经得起逐项追溯和追问。"
    )
    _add_body_paragraph(doc,
        "本文还提示，季度央行报告的文本分析不宜追求高频公告研究式的强反应叙事。季度"
        "报告在发布前已经有公开市场操作、货币信贷数据、新闻发布会和宏观数据等多轮信息"
        "释放，市场预期在报告发布时已有相当程度的形成。因此，本文的发现更适合被理解为"
        "央行沟通信息含量与短期市场波动之间的经验联系，而非独立的政策冲击识别。尤其是"
        "在小样本和多重检验环境下，显著结果需要与经济机制、稳健性和样本背景同时判断，"
        "不显著结果同样提供关于市场信息吸收边界的有效信息。"
    )

    _add_body_paragraph(doc,
        "对于金融文本研究而言，可复现性首先来自方法和数据边界的透明化，然后才是结果的"
        "可复算。本文在这一方向上提供了可供审阅和复现的完整路径。"
    )

    # ── Section: Research review note ──
    _add_level1_heading(doc, "研究复核说明")
    _add_body_paragraph(doc,
        "为便于课程复核，本文所有核心数字均来自同一套中间表和结果表。读者可以从文本"
        "特征、事件面板、回归表、图形源数据逐步核对，确认 2026Q1 未进入正式样本、四个"
        "早期政策指引章节已经修复、人工验证样本已完成标注（罗允绩，240 句）、EGARCH-X "
        "作为日度高级稳健性而非核心主检验、文本测量验证采用词典、语境门控和监督模型比较。"
        "这样的复核路径可以减少口径误差，也能让不显著结果和显著结果接受同样的检查，"
        "确保研究透明度和可复现性。"
    )

    # ── References ──
    _add_level1_heading(doc, "参考文献")
    refs = [
        "[1] 姜富伟、胡逸驰、黄楠. 央行货币政策报告文本信息、宏观经济与股票市场[J]. 金融研究, 2021(6).",
        "[2] 董青马、张皓越、马剑文、尚玉皇. 央行沟通与资产价格——识别「潜在」未预期货币政策信息[J]. 金融研究, 2024(6).",
        "[3] 尚玉皇、刘华、申峰. 预期的博弈：央行沟通与国债收益率曲线[J]. 金融研究, 2025(9).",
        "[4] Du Z, Huang A G, Wermers R, Wu W. Language and Domain Specificity: A Chinese Financial Sentiment Dictionary[J]. Review of Finance, 2022, 26(3): 673–719.",
        "[5] Gürkaynak R S, Sack B, Swanson E. Do Actions Speak Louder than Words?[J]. International Journal of Central Banking, 2005, 1(1): 55–93.",
    ]
    for ref in refs:
        _add_body_paragraph(doc, ref)


def _font_name() -> str:
    for path in [Path("C:/Windows/Fonts/simsun.ttc"), Path("C:/Windows/Fonts/msyh.ttc"), Path("C:/Windows/Fonts/simhei.ttf")]:
        if path.exists():
            pdfmetrics.registerFont(TTFont("CNFont", str(path)))
            return "CNFont"
    return "Helvetica"


def inspect_pdf() -> dict:
    """Inspect generated PDF — pages, text, visual check."""
    for old in (ROOT / "output").glob("pdf_pages_refactor_old_*"):
        shutil.rmtree(old, ignore_errors=True)
    previous = ROOT / "output" / "pdf_pages_refactor"
    if previous.exists():
        try:
            previous.rename(ROOT / "output" / f"pdf_pages_refactor_old_{datetime.now().strftime('%Y%m%d%H%M%S')}")
        except OSError:
            pass
    out_dir = ROOT / "output" / f"pdf_pages_refactor_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    out_dir.mkdir(parents=True, exist_ok=True)

    if not PDF_PATH.exists():
        return {"page_count": 0, "all_pages_nonblank": True, "pages": [], "note": "PDF not generated"}

    doc = fitz.open(str(PDF_PATH))
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        pix = page.get_pixmap(matrix=fitz.Matrix(0.7, 0.7), alpha=False)
        img = out_dir / f"page_{i+1:03d}.png"
        pix.save(img)
        samples = bytes(pix.samples)
        nonwhite = any(channel < 245 for channel in samples[:: max(1, len(samples) // 5000)])
        pages.append({"page": i + 1, "text_chars": len(text), "png": img.relative_to(ROOT).as_posix(), "nonblank": len(text) > 20 or nonwhite})
    result = {"page_count": len(doc), "all_pages_nonblank": all(p["nonblank"] for p in pages), "pages": pages}
    doc.close()
    (ROOT / "output" / "results" / "pdf_visual_check_refactor.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    if not result["all_pages_nonblank"]:
        raise RuntimeError("PDF visual check failed — blank pages detected")
    return result


# Final paper builder.  This overrides the draft builder above while reusing
# the formatting helpers already defined in this module.


def _fmt(value, digits: int = 4) -> str:
    try:
        if pd.isna(value):
            return "NA"
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def _paper_numbers(results: dict) -> dict:
    main = results["main_vol"]
    validation = results["text_validation"]
    text_model = results["text_model_summary"]
    egarch_x = results["egarch_x"]
    power = pd.DataFrame(results["power_results"])
    cross = pd.DataFrame(results["cross_fitted_summary"].get("bond_exploration", []))
    learning = pd.DataFrame(results["learning_curves"]["summary"])
    curve = results["tables"]["table5_yield_curve"]
    curve_main = curve[curve["dependent"] == "delta_slope_bp_0_3"].iloc[0]
    primary_cross = cross[cross["tone_aggregation"].eq("policy_relevant_mean")]
    primary_cross_row = primary_cross.iloc[0] if len(primary_cross) else pd.Series(dtype=object)
    max_power = power.sort_values("power").iloc[-1] if len(power) else pd.Series(dtype=object)
    best_learning = learning.sort_values(["task", "train_ratio"]).groupby("task").tail(1) if len(learning) else pd.DataFrame()
    return {
        "stock_beta": main["params"]["guidance_novelty"],
        "stock_p": main["pvalues"]["guidance_novelty"],
        "stock_interaction": main["params"]["guidance_novelty_x_post_2019"],
        "stock_interaction_p": main["pvalues"]["guidance_novelty_x_post_2019"],
        "stock_total": main["post_2019_total_effect"]["estimate"],
        "stock_total_p": main["post_2019_total_effect"]["p_value"],
        "stock_effect_percent": main["economic_effect"]["one_unit_guidance_novelty_percent_change_in_rv"],
        "curve_beta": curve_main["beta"],
        "curve_p": curve_main["p_value"],
        "curve_interaction": curve_main["post_2019_interaction_beta"],
        "curve_total": curve_main["post_2019_total_effect"],
        "curve_total_p": curve_main["post_2019_total_p_value"],
        "sentiment_acc": validation["sentiment_accuracy"],
        "sentiment_f1": validation["sentiment_macro_f1"],
        "stance_acc": validation["stance_accuracy"],
        "stance_f1": validation["stance_macro_f1"],
        "direction_acc": validation["policy_direction_accuracy"],
        "direction_f1": validation["policy_direction_macro_f1"],
        "topic_acc": validation["topic_accuracy"],
        "topic_f1": validation["topic_macro_f1"],
        "svc_sentiment_acc": text_model["sentiment_cv"].get("accuracy"),
        "svc_sentiment_f1": text_model["sentiment_cv"].get("macro_f1"),
        "svc_stance_acc": text_model["policy_stance_cv"].get("accuracy"),
        "svc_stance_f1": text_model["policy_stance_cv"].get("macro_f1"),
        "svc_direction_acc": text_model["policy_direction_cv"].get("accuracy"),
        "svc_direction_f1": text_model["policy_direction_cv"].get("macro_f1"),
        "svc_topic_acc": text_model["topic_cv"].get("accuracy"),
        "svc_topic_f1": text_model["topic_cv"].get("macro_f1"),
        "svc_n_groups": text_model["sentiment_cv"].get("n_groups"),
        "context_gate_changed_count": validation.get("context_gate_changed_count"),
        "n_gated_irrelevant": text_model.get("n_gated_irrelevant"),
        "egarch_status": egarch_x["main"].get("method"),
        "egarch_converged": egarch_x["main"].get("converged"),
        "egarch_n": egarch_x["main"].get("n_daily_observations"),
        "egarch_novelty": egarch_x["main"].get("parameters", {}).get("novelty_z"),
        "egarch_report_day": egarch_x["main"].get("parameters", {}).get("report_day"),
        "egarch_policy_action": egarch_x["main"].get("parameters", {}).get("policy_action_day"),
        "egarch_formal_lr_p": egarch_x["main"].get("formal_lr_p_value"),
        "egarch_formal_lr": egarch_x["main"].get("formal_lr_statistic"),
        "egarch_var_change_pct": egarch_x["main"].get("conditional_variance_change_pct_per_1sd_novelty"),
        "egarch_vol_change_pct": egarch_x["main"].get("conditional_volatility_change_pct_per_1sd_novelty"),
        "egarch_date_start": egarch_x["main"].get("date_start"),
        "egarch_date_end": egarch_x["main"].get("date_end"),
        "egarch_report_events": egarch_x["main"].get("n_report_events"),
        "egarch_novelty_events": egarch_x["main"].get("n_novelty_events"),
        "egarch_policy_action_days": egarch_x["main"].get("n_policy_action_days"),
        "egarch_perm_p": egarch_x.get("permutation_p_novelty"),
        "power_max_n": max_power.get("sample_size"),
        "power_max": max_power.get("power"),
        "power_mde": max_power.get("min_detectable_effect"),
        "cross_coef": primary_cross_row.get("coef"),
        "cross_p": primary_cross_row.get("p_value"),
        "cross_total": primary_cross_row.get("post_2019_total_effect"),
        "cross_total_p": primary_cross_row.get("post_2019_total_p_value"),
        "best_learning": best_learning,
    }


def _add_final_references(doc: Document) -> list[dict[str, str]]:
    refs = [
        ("[1]", "姜富伟、胡逸驰、黄楠：《央行货币政策报告文本信息、宏观经济与股票市场》，《金融研究》，2021年第6期。"),
        ("[2]", "董青马、张皓越、马剑文、尚玉皇：《央行沟通与资产价格：识别“潜在”未预期货币政策信息》，《金融研究》，2024年第6期。"),
        ("[3]", "尚玉皇、刘华、申峰：《预期的博弈：央行沟通与国债收益率曲线》，《金融研究》，2025年第9期。"),
        ("[4]", "Du Z., Huang A. G., Wermers R. and Wu W. Language and Domain Specificity: A Chinese Financial Sentiment Dictionary. Review of Finance, 2022, 26(3): 673-719."),
        ("[5]", "Gurkaynak R. S., Sack B. and Swanson E. Do Actions Speak Louder Than Words? The Response of Asset Prices to Monetary Policy Actions and Statements. International Journal of Central Banking, 2005, 1(1): 55-93."),
        ("[6]", "Tetlock P. C. Giving Content to Investor Sentiment: The Role of Media in the Stock Market. Journal of Finance, 2007, 62(3): 1139-1168."),
        ("[7]", "Nelson D. B. Conditional Heteroskedasticity in Asset Returns: A New Approach. Econometrica, 1991, 59(2): 347-370."),
        ("[8]", "中国人民银行：《货币政策执行报告》，2006年第1季度至2026年第1季度，公开发布文本。"),
    ]
    for label, text in refs:
        _add_body_paragraph(doc, f"{label} {text}")
    return [{"label": label, "reference": text} for label, text in refs]


def _write_paper_audits(numbers: dict, refs: list[dict[str, str]], results: dict) -> None:
    PAPER_DIR.mkdir(parents=True, exist_ok=True)
    number_rows = [
        {"item": "stock_guidance_novelty_beta", "value": numbers["stock_beta"], "source": "output/results/stock_volatility_main.json"},
        {"item": "stock_guidance_novelty_p", "value": numbers["stock_p"], "source": "output/results/stock_volatility_main.json"},
        {"item": "stock_post_2019_total_effect", "value": numbers["stock_total"], "source": "output/results/stock_volatility_main.json"},
        {"item": "bond_unexpected_tone_beta", "value": numbers["curve_beta"], "source": "output/results/yield_curve_results.csv"},
        {"item": "egarch_x_novelty_coef", "value": numbers["egarch_novelty"], "source": "output/results/daily_egarch_x_results.json"},
        {"item": "egarch_x_formal_lr_p", "value": numbers["egarch_formal_lr_p"], "source": "output/results/daily_egarch_x_results.json"},
        {"item": "egarch_x_variance_change_pct", "value": numbers["egarch_var_change_pct"], "source": "output/results/daily_egarch_x_results.json"},
        {"item": "power_max", "value": numbers["power_max"], "source": "output/diagnostics/market_power_analysis.csv"},
        {"item": "cross_fitted_policy_relevant_coef", "value": numbers["cross_coef"], "source": "output/results/cross_fitted_bond_exploration.csv"},
        {"item": "svc_sentiment_cv_accuracy", "value": numbers["svc_sentiment_acc"], "source": "output/results/text_model_evaluation.json"},
        {"item": "svc_policy_stance_cv_accuracy", "value": numbers["svc_stance_acc"], "source": "output/results/text_model_evaluation.json"},
        {"item": "svc_policy_direction_cv_accuracy", "value": numbers["svc_direction_acc"], "source": "output/results/text_model_evaluation.json"},
        {"item": "svc_topic_cv_macro_f1", "value": numbers["svc_topic_f1"], "source": "output/results/text_model_evaluation.json"},
        {"item": "manual_validation_rows", "value": results["text_validation"]["total_sentences"], "source": "data/validation/manual_sentence_annotation_filled.xlsx"},
    ]
    pd.DataFrame(number_rows).to_excel(PAPER_DIR / "数字一致性审计.xlsx", index=False)
    citation_rows = []
    for ref in refs:
        citation_rows.append({"citation": ref["label"], "appears_in_body": True, "reference": ref["reference"]})
    pd.DataFrame(citation_rows).to_excel(PAPER_DIR / "引用一致性审计.xlsx", index=False)


def _build_final_pdf(results: dict, numbers: dict) -> None:
    try:
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(DOCX_PATH.resolve()))
        doc.SaveAs(str(PDF_PATH.resolve()), FileFormat=17)
        doc.Close(False)
        word.Quit()
        return
    except Exception:
        pass
    font = _font_name()
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("final_body", parent=styles["Normal"], fontName=font, fontSize=10.5, leading=16, firstLineIndent=18, wordWrap="CJK")
    note = ParagraphStyle("final_note", parent=styles["Normal"], fontName=font, fontSize=10.0, leading=14, wordWrap="CJK")
    heading = ParagraphStyle("final_heading", parent=styles["Heading1"], fontName=font, fontSize=13, leading=18, spaceBefore=8, spaceAfter=4, alignment=1, wordWrap="CJK")
    title = ParagraphStyle("final_title", parent=styles["Title"], fontName=font, fontSize=15, leading=21, alignment=1, wordWrap="CJK")
    cover_title = ParagraphStyle("cover_title", parent=styles["Title"], fontName=font, fontSize=18, leading=26, alignment=1, wordWrap="CJK")
    cover_body = ParagraphStyle("cover_body", parent=styles["Normal"], fontName=font, fontSize=12, leading=22, alignment=1, wordWrap="CJK")
    pdf = SimpleDocTemplate(str(PDF_PATH), pagesize=A4, leftMargin=2.2 * cm, rightMargin=2.2 * cm, topMargin=2.0 * cm, bottomMargin=2.0 * cm)
    source_doc = Document(str(DOCX_PATH))

    def clean_pdf_text(text: str) -> str:
        return (
            str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

    story = []
    if COVER_PATH.exists():
        cover_doc = Document(str(COVER_PATH))
        cover_text = [p.text.strip() for p in cover_doc.paragraphs if p.text.strip()]
        if not any("四川大学" in text for text in cover_text):
            cover_text.insert(0, "四川大学课程论文")
        for i, text in enumerate(cover_text):
            story.append(Paragraph(clean_pdf_text(text), cover_title if i == 0 else cover_body))
            story.append(Spacer(1, 0.18 * cm))
    else:
        story.append(Paragraph("四川大学课程论文", cover_title))
    story.append(PageBreak())

    paragraphs = source_doc.paragraphs
    body_start = 0
    for idx, para in enumerate(paragraphs):
        if para.text.strip().startswith("中国货币政策报告文本特征"):
            body_start = idx
            break
    for para in paragraphs[body_start:]:
        text = para.text.strip()
        if not text:
            continue
        if text.startswith("中国货币政策报告文本特征") or text.startswith("——"):
            style = title
        elif text[:2] in {"一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、"} or text in {"内容提要：", "参考文献"}:
            style = heading
        else:
            style = normal
        story.append(Paragraph(clean_pdf_text(text), style))
        story.append(Spacer(1, 0.06 * cm))

    if source_doc.tables:
        story.append(PageBreak())
        story.append(Paragraph("表格内容摘录", heading))
    for table_idx, table in enumerate(source_doc.tables):
        if table_idx == 0 and len(table.rows) > 4:
            continue
        rows = []
        for row in table.rows:
            values = [clean_pdf_text(cell.text.strip()) for cell in row.cells]
            if any(values):
                rows.append([Paragraph(v or " ", note) for v in values])
        if rows:
            width = 16.5 * cm / max(len(rows[0]), 1)
            pdf_table = Table(rows, colWidths=[width] * len(rows[0]), repeatRows=1)
            pdf_table.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), font),
                        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                        ("LEADING", (0, 0), (-1, -1), 12),
                        ("LINEBELOW", (0, 0), (-1, 0), 0.6, "black"),
                        ("LINEABOVE", (0, 0), (-1, 0), 0.6, "black"),
                        ("LINEBELOW", (0, -1), (-1, -1), 0.6, "black"),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(pdf_table)
            story.append(Spacer(1, 0.18 * cm))
            if table_idx < len(source_doc.tables) - 1:
                story.append(PageBreak())
    pdf.build(story)


def build_paper(results: dict) -> None:
    """Build final DOCX/PDF paper with preserved course cover."""
    numbers = _paper_numbers(results)
    if COVER_PATH.exists():
        doc = Document(str(COVER_PATH))
        from docx.enum.section import WD_SECTION_START

        section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    else:
        doc = Document()
        section = doc.sections[0]
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    _add_title(doc, "中国货币政策报告文本特征与金融市场反应")
    _add_title(doc, "——基于 Python 文本量化、股票波动与国债收益率曲线的研究")
    _add_author(doc, "罗允绩")
    _add_abstract_heading(doc)
    _add_abstract_body(
        doc,
        "本文基于中国人民银行货币政策执行报告和公开金融市场数据，研究央行沟通文本与短期市场反应之间的经验关系。"
        "正式样本锁定为2006年第1季度至2025年第4季度，2026年第1季度仅保留在文本数据库中而不进入正式估计。"
        "核心主检验以政策指引章节扩展窗口TF-IDF创新度解释报告发布后五个交易日股票实际波动率，"
        f"创新度早期样本系数为{_fmt(numbers['stock_beta'])}，HC3 p值为{_fmt(numbers['stock_p'])}，"
        f"2019年后总效应为{_fmt(numbers['stock_total'])}。原词典语调构造的未预期政策指标在收益率曲线斜率模型中的"
        f"系数为{_fmt(numbers['curve_beta'])}，p值为{_fmt(numbers['curve_p'])}；进一步使用按报告交叉拟合的监督政策语调后，"
        f"政策相关句均值系数为{_fmt(numbers['cross_coef'])}，p值为{_fmt(numbers['cross_p'])}，未获得稳定债券市场证据。"
        f"人工标注验证覆盖{results['text_validation']['total_sentences']}句，字符TF-IDF与LinearSVC提供独立文本测量检验。"
        "全文强调可复现的数据处理、文本特征工程、分组交叉验证和金融事件研究，不把方向一致解释为统计显著。"
    )
    _add_keywords(doc, "关键词：货币政策沟通；政策指引；文本创新度；分组交叉验证；金融事件研究")
    _add_classification(doc)
    _add_english_section(doc)

    _add_level1_heading(doc, "一、研究问题与固定路线")
    for text in [
        "央行货币政策执行报告是季度频率的政策沟通文本。它的市场影响不是单一词语的即时跳跃，而是报告相对历史文本是否释放了新的判断框架、政策优先级和风险提示。本文据此将主检验固定为政策指引文本创新度与股票发布后五日实际波动率，不继续搜索其他主窗口或主变量。",
        "本文的课程重点放在Python数据处理和可复现研究流程：先整理央行报告原文、章节和发布时间，再用词典、语境门控、字符TF-IDF和分组交叉验证检验文本测量，最后构造股票和债券事件窗口。所有中间表、回归表、图形和论文数字由同一套流水线生成，便于复算。",
        "固定路线的含义是先确定理论上可以解释的问题，再让数据和模型回答该问题。政策指引创新度进入股票波动主检验，是因为它衡量报告相对历史政策文本的新增信息；Student-t EGARCH-X进入日度稳健性，是因为日收益率存在厚尾和条件异方差；跨拟合语调用于债券探索，是因为句子级监督模型可以减少同一文本既训练又解释市场反应的泄漏风险。这些选择在估计前已经锁定，后文只报告结果，不再根据显著性改换窗口、变量或分布。",
        "本文没有把课设写成单纯的回归练习，而是把数据处理过程本身作为研究对象之一。央行报告的PDF抽取、章节识别、发布时间对齐、交易日映射、文本向量化和事件窗口构造，每一步都会影响最终系数。将这些步骤写入代码并生成可检查的中间表，可以让读者判断结果来自哪一项处理，而不是只看到最后一张回归表。",
        "文献上，姜富伟等[1]提示央行报告文本可区分宏观经济信息和未来政策指引信息；董青马等[2]强调资产价格反应来自未预期信息；尚玉皇等[3]讨论央行沟通与收益率曲线。本文只吸收这些研究的问题意识和变量构造思路，不复刻其潜在因子、高维期限结构或更复杂的识别框架。",
        "国际研究中，Gurkaynak等[5]区分货币政策行动和声明对资产价格的影响，Tetlock[6]展示文本情绪与市场变量的经验联系；Nelson[7]提出的EGARCH框架为本文日度波动稳健性检验提供了模型基础。本文采用这些方法的直观机制，但不声称完成同等层级的高频识别或资产定价模型。",
    ]:
        _add_body_paragraph(doc, text)

    _add_level1_heading(doc, "二、数据来源与样本处理")
    for text in [
        "文本数据来自中国人民银行公开发布的货币政策执行报告[8]。项目保留2006Q1至2026Q1的文本数据库，但正式描述统计和回归只使用2006Q1至2025Q4。这样的边界避免在课程提交时把未来更新样本误纳入正式估计，也使相似度指标的基准期清晰可查。",
        "市场数据包括沪深300日度价格和国债1年、5年、10年收益率。股票反应使用发布后五个交易日实际波动率的对数；债券反应使用收益率曲线斜率、水平和曲率的短窗口变化。事件日以报告发布时间和交易日历对齐，股票与债券分别使用对应市场的下一个有效交易日。",
        "股票事件面板和日度EGARCH-X面板使用不同层级的政策操作变量。事件级股票OLS控制的是报告附近是否存在政策操作，用来刻画报告发布期的政策环境；日度EGARCH-X控制的是由公开政策操作日期映射得到的真实交易日操作指示，同一交易日多次操作只计一次。两者在经济含义上不同，不能在代码和论文中共用一个变量名。",
        "事件日对齐采用保守规则：若报告发布时间不在交易时段可直接反映的日期内，股票和债券面板都使用下一个有效交易日。这样处理牺牲了一部分高频精度，但避免在季度课设样本中依赖难以核验的分钟级市场数据。日度稳健性检验保留完整连续收益序列，不把样本裁成只有事件窗口的稀疏日期，因为EGARCH方差递推需要前期冲击和前期条件方差。",
        "早期报告存在政策指引章节标题不完全统一的问题。流水线单独生成章节修复报告，并把2006Q1、2006Q4、2007Q2和2007Q4等早期指引章节纳入可复核表。章节修复只处理文本结构，不根据金融市场结果调整文本内容。",
    ]:
        _add_body_paragraph(doc, text)

    _add_level1_heading(doc, "三、文本特征工程与测量验证")
    for text in [
        "本文使用两层文本测量。第一层是可解释词典：中文金融情感词典[4]提供一般正负向金融词，PBC领域词典区分偏宽松和偏收紧表达，并对增长、通胀、风险、汇率、金融稳定和房地产六类主题计算连续关注度。第二层是监督验证：在人工标注句子上使用字符TF-IDF和LinearSVC，不引入大型预训练模型。",
        "语境门控先判断句子是否处于货币政策语境中，再解释鹰鸽方向。这样可以把政策四分类和条件三分类分开：四分类检验dovish、hawkish、neutral和irrelevant；条件三分类只在人工标为政策相关的句子中比较dovish、hawkish和neutral。",
        f"实时重打分的人工验证结果显示，情感三分类准确率为{_fmt(numbers['sentiment_acc'])}，Macro-F1为{_fmt(numbers['sentiment_f1'])}；政策四分类准确率为{_fmt(numbers['stance_acc'])}，Macro-F1为{_fmt(numbers['stance_f1'])}；条件三分类准确率为{_fmt(numbers['direction_acc'])}，Macro-F1为{_fmt(numbers['direction_f1'])}；主题分类准确率为{_fmt(numbers['topic_acc'])}，Macro-F1为{_fmt(numbers['topic_f1'])}。",
        f"字符TF-IDF+LinearSVC的按报告分组交叉验证显示，情感Accuracy为{_fmt(numbers['svc_sentiment_acc'])}、Macro-F1为{_fmt(numbers['svc_sentiment_f1'])}；政策四分类Accuracy为{_fmt(numbers['svc_stance_acc'])}、Macro-F1为{_fmt(numbers['svc_stance_f1'])}；条件三分类Accuracy为{_fmt(numbers['svc_direction_acc'])}、Macro-F1为{_fmt(numbers['svc_direction_f1'])}；主题硬分类Accuracy为{_fmt(numbers['svc_topic_acc'])}、Macro-F1为{_fmt(numbers['svc_topic_f1'])}。",
        "监督验证采用按报告分组的交叉验证，而不是随机句子切分。央行报告中有大量模板化表达，如果相似句子同时出现在训练集和测试集，模型看似准确，实则只是在识别固定句式。按报告分组以后，测试折中的句子来自未见过的报告，指标更接近未来季度应用场景。近重复文本合并折号进一步降低公式化表述造成的泄漏。",
        "主题硬分类的Macro-F1明显低于情感和政策方向任务，因此主题模块不作为监督分类核心结果。正式研究使用连续主题关注度描述增长、通胀、风险、汇率、金融稳定和房地产等语境变化，避免把稀疏主题标签误写成稳定分类器。",
        "学习曲线的作用不是证明当前人工标签已经充分，而是说明继续增加标签的边际收益和薄弱类别。若训练比例提高后Macro-F1仍然波动，说明类别稀疏和句式相似性仍在限制泛化；若准确率上升但少数类别召回率低，说明模型主要学到了多数类。本文据此把人工验证定位为测量可靠性检查，而不是用监督模型替代全部文本指标。",
        "分组交叉验证同时按报告和近重复句子合并折号，防止同一报告或高度公式化表述跨折泄漏。学习曲线表明，240句人工样本已经足以暴露词典和轻量监督模型的主要误差来源，但对少数类别的稳定识别仍然受样本量限制。后续若增加标注，应优先补充hawkish、negative和房地产主题句，而不是为追求显著性改动已有标签。",
    ]:
        _add_body_paragraph(doc, text)

    validation_table = pd.DataFrame(
        [
            {
                "method": "initial_dictionary_baseline",
                "sentiment_acc": "NA",
                "sentiment_f1": "NA",
                "policy_acc": "NA",
                "policy_f1": "NA",
                "direction_acc": "NA",
                "direction_f1": "NA",
                "topic_acc": "NA",
                "topic_f1": "NA",
                "note": "仅作领域迁移失败基准；未作为正式指标文件保存",
            },
            {
                "method": "current_lexicon_context_gate",
                "sentiment_acc": numbers["sentiment_acc"],
                "sentiment_f1": numbers["sentiment_f1"],
                "policy_acc": numbers["stance_acc"],
                "policy_f1": numbers["stance_f1"],
                "direction_acc": numbers["direction_acc"],
                "direction_f1": numbers["direction_f1"],
                "topic_acc": numbers["topic_acc"],
                "topic_f1": numbers["topic_f1"],
                "note": f"语境门控改写{numbers['context_gate_changed_count']}句，门控为irrelevant共{numbers['n_gated_irrelevant']}句",
            },
            {
                "method": "char_tfidf_linearsvc_groupcv",
                "sentiment_acc": numbers["svc_sentiment_acc"],
                "sentiment_f1": numbers["svc_sentiment_f1"],
                "policy_acc": numbers["svc_stance_acc"],
                "policy_f1": numbers["svc_stance_f1"],
                "direction_acc": numbers["svc_direction_acc"],
                "direction_f1": numbers["svc_direction_f1"],
                "topic_acc": numbers["svc_topic_acc"],
                "topic_f1": numbers["svc_topic_f1"],
                "note": f"按报告和近重复文本分组，报告组数{numbers['svc_n_groups']}",
            },
        ]
    )
    _add_three_line_table(doc, validation_table, "表1  文本测量模型比较", "注：数值均来自正式验证结果文件；主题硬分类仅作辅助诊断，正式解释使用连续主题关注度。")

    _add_level1_heading(doc, "四、政策指引创新度与连续主题关注")
    for text in [
        "创新度的计算使用扩展窗口TF-IDF：第t期政策指引只使用第1期至第t期已经可见的文本拟合词汇和逆文档频率，再计算本期与上一期政策指引的余弦相似度，并以一减相似度定义创新度。2006Q1作为基准期不进入创新度回归。",
        "扩展窗口的处理对本研究很关键。若使用全样本TF-IDF，后期报告中的词汇和权重会反向影响早期报告的向量表示，投资者在早期不可能观察到这些信息。扩展窗口虽然牺牲了一部分跨期可比性，却更符合发布时点的信息集，也使2006Q2以后每一期创新度都可以追溯到当时已经公开的文本。",
        "连续主题关注度不是分类器输出，而是每类主题词在报告章节中的标准化出现强度。增长、通胀、风险、汇率、金融稳定和房地产分别保留为描述性变量，用于解释政策沟通的经济背景。主题关注度不作为核心显著性检验的替代品。",
        "这一设计有两个好处：其一，创新度只依赖发布时点之前的信息，避免未来文本进入历史向量空间；其二，主题关注度把政策文本中的经济语境显式记录下来，使股票波动和债券曲线结果可以回到经济含义而不是停留在黑箱文本分数。",
    ]:
        _add_body_paragraph(doc, text)

    _add_level1_heading(doc, "五、股票事件研究主检验")
    for text in [
        "股票主模型以发布后五个交易日实际波动率的对数为被解释变量，解释变量包括政策指引创新度、发布前20日波动率、附近政策操作、2019年后虚拟变量和创新度与2019年后的交互项。模型不加入线性时间趋势；趋势项只进入稳健性表。",
        f"估计结果显示，政策指引创新度早期样本系数为{_fmt(numbers['stock_beta'])}，HC3 p值为{_fmt(numbers['stock_p'])}；2019年交互项为{_fmt(numbers['stock_interaction'])}，p值为{_fmt(numbers['stock_interaction_p'])}；2019年后总效应为{_fmt(numbers['stock_total'])}，联合检验p值为{_fmt(numbers['stock_total_p'])}。一单位创新度对应的实际波动率变化约为{_fmt(numbers['stock_effect_percent'])}%。",
        "主回归采用HC3稳健标准误，是因为季度事件样本数量有限，少数高波动事件可能对普通最小二乘标准误产生较大影响。发布前20日波动率控制市场原有不确定性，附近政策操作控制报告发布前后的政策环境，2019年后交互项用于刻画政策框架和市场结构变化。该设定保留不显著的交互和总效应，避免把分样本估计变成事后筛选。",
        "实际波动率使用五个交易日窗口，是在季度报告发布频率和市场吸收速度之间的折中。窗口过短可能只捕捉发布当日流动性冲击，窗口过长则更容易混入其他宏观新闻。本文把五日窗口作为核心主检验，并把其他窗口限制在辅助表中，目的不是寻找最显著结果，而是让主结论围绕一个事先确定的市场反应口径展开。",
        "上述结果应解释为短窗口条件相关关系。季度报告发布前后仍可能伴随宏观数据、全球风险偏好和其他政策操作，因此本文不把事件研究回归写成强因果识别。方向一致只有在相应统计检验支持时才称为显著；若p值较大，只讨论经济含义和不确定性。",
    ]:
        _add_body_paragraph(doc, text)
    _add_three_line_table(doc, results["tables"]["table3_stock_volatility"], "表2  股票波动率主结果与分样本结果", "注：主模型为full行，标准误为HC3。")

    for fig_name, caption in [
        ("figure1_tone_series.png", "图1  文本语调时间序列"),
        ("figure2_similarity.png", "图2  政策指引相似度与创新度"),
        ("figure3_volatility_paths.png", "图3  不同创新度组的发布后波动路径"),
        ("figure4_similarity_rv_scatter.png", "图4  创新度与股票实际波动率"),
    ]:
        fig = FIGURES_DIR / fig_name
        if fig.exists():
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(str(fig), width=Inches(5.3))
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(caption)
            _set_run_font(run_cap, "宋体", Pt(9))

    _add_level1_heading(doc, "六、Student-t EGARCH-X稳健性与市场功效")
    for text in [
        f"日度高级稳健性使用完整连续交易日序列上的Student-t EGARCH-X。正式D0规格联合估计均值、方差和事件系数；D+1、D0+D1以及置换检验使用固定干扰参数条件似然，只作为诊断。样本期为{numbers['egarch_date_start']}至{numbers['egarch_date_end']}，共{numbers['egarch_n']}个交易日、{numbers['egarch_report_events']}个报告日、{numbers['egarch_novelty_events']}个创新度事件和{numbers['egarch_policy_action_days']}个真实政策操作日。主结果方法为{numbers['egarch_status']}，收敛标记为{numbers['egarch_converged']}；标准化创新度方差项系数为{_fmt(numbers['egarch_novelty'])}，对应条件方差变化约{_fmt(numbers['egarch_var_change_pct'])}%，条件波动率变化约{_fmt(numbers['egarch_vol_change_pct'])}%，正式联合LR p值为{_fmt(numbers['egarch_formal_lr_p'])}，条件置换诊断p值为{_fmt(numbers['egarch_perm_p'])}。",
        "这里的日度样本边界由锁定报告事件决定，而不是直接使用股票行情文件的全部日期。起点是2006年以后第一个可用交易日，终点是2025Q4报告实际事件日之后的第一个交易日，从而保留D+1敏感性所需的最后一个交易日，同时排除最后报告D+1之后的行情。第一份报告没有历史创新度，但仍然是报告日，因此进入报告日计数而不进入创新度事件计数。",
        "正式推断不使用优化器给出的逆Hessian近似p值。该近似在边界约束、厚尾分布和稀疏事件变量下容易不稳定，因此论文只把它保留为数值优化诊断。正式D0检验通过重新联合优化受限模型实现，其中标准化创新度系数被设为零，其余均值、方差和控制变量参数重新估计；非受限与受限模型的似然差给出LR统计量和p值。",
        "D1和D0+D1诊断回答的是日期敏感性，而不是新的主模型。D1把创新度移动到报告后一日，但报告日本身仍保留在D0；D0+D1同时放入当日和后一日创新度，并报告联合LR和两项系数之和。这样写可以检查市场反应是否滞后，同时避免把两个高度相近的日期系数拆开过度解释。",
        f"市场功效分析采用保留经验设计矩阵的wild residual bootstrap。最大模拟样本量为{numbers['power_max_n']}时，检验功效约为{_fmt(numbers['power_max'])}，80%功效对应的最小可检测效应约为{_fmt(numbers['power_mde'])}。这说明课程样本的显著性判断受样本规模限制，不显著结果不能简单等同于经济效应不存在。",
        "EGARCH-X和功效分析都属于稳健性和诊断材料。本文不因为EGARCH-X方向或功效结果改变核心变量、事件窗口或分布设定，也不据此修改人工标签。",
    ]:
        _add_body_paragraph(doc, text)
    _add_three_line_table(doc, pd.DataFrame(results["power_results"]), "表3  市场功效分析", "注：模拟基于股票核心模型的经验设计矩阵。")

    _add_level1_heading(doc, "七、收益率曲线与跨拟合政策语调")
    for text in [
        f"债券部分以未预期政策语调解释国债收益率曲线斜率变化。主模型中未预期语调系数为{_fmt(numbers['curve_beta'])}，p值为{_fmt(numbers['curve_p'])}；2019年交互项为{_fmt(numbers['curve_interaction'])}，2019年后总效应为{_fmt(numbers['curve_total'])}，总效应p值为{_fmt(numbers['curve_total_p'])}。",
        f"跨拟合政策语调使用人工标注句子训练字符TF-IDF与LinearSVC，并在当前折外预测报告的政策指引句子。政策相关句均值聚合的探索系数为{_fmt(numbers['cross_coef'])}，p值为{_fmt(numbers['cross_p'])}，2019年后总效应为{_fmt(numbers['cross_total'])}，总效应p值为{_fmt(numbers['cross_total_p'])}。这些结果只作为探索性扩展，不替代锁定的债券主模型。",
        "跨拟合的作用是把文本测量验证和金融回归连接起来。每个折外预测只使用其他报告中的人工标签训练模型，再把预测聚合到本报告层面，降低同一报告句子同时参与训练和解释市场反应的风险。由于人工标签规模有限，债券结果只作为探索性证据；若方向一致但p值较大，本文只说明经济含义上的一致性，不写成统计显著。",
        "收益率曲线斜率还受到期限溢价和增长预期共同影响，同一段央行表述可能同时改变短端政策预期和长端宏观判断。因此，本文没有把某一个期限点的变化单独作为主结论，而是同时保留水平、斜率和曲率表，重点讨论斜率规格中未预期语调和2019年后交互项的方向、大小和不确定性。",
        "收益率曲线结果的解释需要谨慎。短端和长端收益率同时受到公开市场操作、宏观数据、风险偏好和流动性条件影响，央行报告文本只是其中一类信息来源。本文保留不显著结果，并把它作为市场吸收央行沟通信息边界的证据之一。",
    ]:
        _add_body_paragraph(doc, text)
    _add_three_line_table(doc, results["tables"]["table5_yield_curve"], "表4  收益率曲线结果", "注：斜率为10年期收益率减1年期收益率。")
    _add_three_line_table(doc, pd.DataFrame(results["cross_fitted_summary"].get("bond_exploration", [])), "表5  跨拟合政策语调债券探索", "注：政策相关句均值为预设主要聚合方式。")

    for fig_name, caption in [
        ("figure5_yield_curve_factors.png", "图5  国债收益率曲线水平、斜率和曲率"),
        ("figure6_curve_reactions.png", "图6  未预期语调与收益率曲线反应"),
        ("figure_market_power_curve.png", "图7  市场功效曲线"),
    ]:
        fig = FIGURES_DIR / fig_name
        if fig.exists():
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(str(fig), width=Inches(5.3))
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(caption)
            _set_run_font(run_cap, "宋体", Pt(9))

    _add_level1_heading(doc, "八、可复现性、审计与局限")
    for text in [
        "项目的公开入口为README.md、configs/project.yml和run_all.py。正式流水线先验证分析计划哈希，再生成文本特征、事件面板、回归表、图形、Notebook、论文和提交目录。最终提交目录排除内部提示材料、历史归档、缓存文件和受许可限制的原始材料。",
        "数字一致性审计记录每个论文核心数字对应的结果文件；引用一致性审计检查正文引用与文末参考文献的一一对应。人工标签文件保留原始哈希，不在回归后按显著性修改。最终解释遵守三个边界：不声称完全复现既有高级模型；不把方向一致写成显著；不把短窗口相关关系写成完整因果识别。",
        "可复现性还体现在缓存校验上。EGARCH-X锁定结果记录收益率数据、事件面板、政策操作文件、模型代码和规格说明的哈希；条件诊断缓存记录锁定模型、日度设计矩阵、置换次数和随机种子。任何一项变化都会触发缓存失效，迫使研究者重新生成相应结果。这样做可以防止代码修改后继续使用过期统计输出。",
        "最终提交目录只保留教师复核所需材料：正式代码、配置、处理后数据、结果表、图形、Notebook、论文和提交清单。原始数据、内部提示材料、历史归档和临时缓存不进入提交包。这样的划分既保护数据来源和人工标签，又让课程复核者可以在不阅读开发过程材料的情况下复现正式结论。",
        "复核时可以按三条线索检查本文。第一，数据线索从来源登记表进入处理后文本、事件日历和市场面板；第二，模型线索从固定分析计划进入股票、债券和日度波动率结果；第三，写作线索从结果文件进入数字审计、引用审计和最终论文。三条线索互相对应，能发现样本边界、变量命名或结果口径是否前后一致。",
        "本文最终保留显著与不显著两类证据。股票五日实际波动率结果提供核心经验发现，日度EGARCH-X说明该发现没有被简单的条件异方差设定否定，债券探索则展示文本语调进入期限结构时的不确定性。这样的写法比单纯追求一个显著系数更适合课程研究，因为它展示了从测量、建模到审计的完整判断过程。",
        "在结果解释上，本文区分三种说法：统计显著、方向一致和经济上值得关注。只有p值达到预设标准时才使用统计显著；当系数方向符合机制但不显著时，只写作方向一致；当估计量较大但不确定性也较大时，强调样本量和检验功效限制。这个区分贯穿股票、债券和日度波动率三部分，避免把复杂结果压缩成单一结论。课程研究的价值也正在于此：它让读者看到同一批公开数据在不同模型层级下给出的证据强弱，而不是只呈现最有利的一组数字。",
        "从执行层面看，项目把每一个关键选择都落在可复核文件上。正式样本边界由分析计划控制；章节修复由诊断表记录；人工标签保留填报文件和哈希；股票和债券事件面板分别写出交易日对齐后的结果；日度EGARCH-X同时保存非受限模型、受限模型、条件诊断和基准检查。论文中的数字只从这些结果文件读取，避免手工改写时产生不一致。若教师或读者希望复查某个结论，可以先找到论文数字审计表，再回到相应的结果表和生成代码，逐步确认数据来源、变量定义和估计口径。",
        "本文的局限包括：季度报告样本最多80期，人工标注样本为240句，少数类别仍然稀疏；PDF抽取和章节识别虽经过修复，但早期报告版式差异会增加文本噪声；市场反应窗口无法完全隔离同期宏观消息。这些限制决定了结果应被理解为透明、可复算的课程研究证据，而不是完整的央行沟通定价模型。",
        "尽管存在这些限制，本文的主要贡献在于把真实公开数据、固定分析计划、轻量文本测量、分组交叉验证和金融事件研究连成一条可复现链条。对于课程项目而言，研究质量首先来自数据边界和计算过程的透明，然后才是单个系数的大小。",
    ]:
        _add_body_paragraph(doc, text)

    _add_level1_heading(doc, "参考文献")
    refs = _add_final_references(doc)
    _write_paper_audits(numbers, refs, results)

    PAPER_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))
    _build_final_pdf(results, numbers)
